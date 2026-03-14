import streamlit as st
import anthropic
import pdfplumber
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io, json, re, time
import requests
from bs4 import BeautifulSoup

MAX_CV_CHARS = 15_000

# ─── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(page_title="CV Optimizer ATS", page_icon="🎯", layout="centered")

st.markdown("""
<style>
.block-container { padding-top: 2rem; padding-bottom: 2rem; }
div[data-testid="stDownloadButton"] button {
    background-color: #1B6CA8; color: white;
    font-size: 1rem; padding: 0.6rem 1.5rem;
    border-radius: 6px; width: 100%; }
.coach-card { background: #F0F9FF; border-left: 4px solid #2E75B6;
              padding: 0.75rem 1rem; border-radius: 4px;
              margin-bottom: 0.5rem; font-size: 0.95rem; }
.score-explain { background: #F8F8F8; border-radius: 8px;
                 padding: 0.75rem 1rem; font-size: 0.9rem;
                 color: #444; margin-top: 0.5rem; }
.warn-box { background: #FFF8E1; border-left: 3px solid #FFA000;
            padding: 0.6rem 0.9rem; border-radius: 4px;
            font-size: 0.88rem; color: #555; margin-bottom: 0.75rem; }
.tip-box  { background: #E8F5E9; border-left: 3px solid #4CAF50;
            padding: 0.6rem 0.9rem; border-radius: 4px;
            font-size: 0.88rem; color: #2E7D32; margin-bottom: 0.75rem; }
.template-preview { border: 1px solid #DDD; border-radius: 6px;
                    padding: 0.6rem 0.8rem; font-size: 0.82rem;
                    background: #FAFAFA; margin-top: 0.3rem;
                    font-family: monospace; line-height: 1.5; }
</style>
""", unsafe_allow_html=True)

# ─── Header ───────────────────────────────────────────────────────────────────
st.title("🎯 CV Optimizer ATS")
st.markdown("Adapta tu CV a cualquier oferta laboral y supera los filtros automáticos.")

# ─── API Key ──────────────────────────────────────────────────────────────────
def get_secret_key():
    try:
        return st.secrets["ANTHROPIC_API_KEY"]
    except (KeyError, FileNotFoundError):
        return None

_secret_key = get_secret_key()

# ─── Sidebar ──────────────────────────────────────────────────────────────────
with st.sidebar:
    st.header("⚙️ Configuración")

    if _secret_key:
        api_key = _secret_key
        if st.session_state.get("api_credits_error"):
            st.warning("Servicio sin saldo. Puedes usar tu propia API Key:")
            user_key = st.text_input("🔑 Tu Anthropic API Key", type="password")
            if user_key:
                api_key = user_key
    else:
        st.warning("⚠️ Ingresa tu API Key")
        api_key = st.text_input("🔑 Anthropic API Key", type="password",
                                 help="Obtén tu key en console.anthropic.com")

    st.markdown("---")
    st.markdown("**📐 Formato del CV**")

    max_pages = st.slider("Páginas máximas", 1, 3, 2,
        help="Claude seleccionará el contenido más relevante para ajustarse a este límite.")

    font_family = st.selectbox("Tipografía",
        ["Calibri", "Arial", "Georgia", "Times New Roman", "Trebuchet MS"], index=0,
        help="Calibri y Arial son las más amigables con los ATS.")

    font_size = st.select_slider("Tamaño de letra",
        options=[9, 10, 10.5, 11, 12], value=10,
        help="Los títulos se ajustan proporcionalmente.")

    if st.session_state.get("cv_data"):
        st.markdown("---")
        st.success("✅ Análisis listo")
        st.markdown("**Cambia template, fuente o páginas** y descarga cuantas versiones quieras — sin volver a consumir créditos.")
        if st.button("🔄 Generar nueva versión DOCX", use_container_width=True):
            st.session_state["regen_docx"] = True

    st.markdown("---")
    st.markdown("**¿Cómo funciona?**")
    st.markdown("1. Sube tu **CV maestro** con toda tu experiencia")
    st.markdown("2. Pega o linkea la oferta laboral")
    st.markdown("3. Elige template y descarga")
    st.markdown("4. Repite para otras ofertas sin límite")
    st.markdown("---")
    st.caption("Powered by Claude AI · Anthropic")

# ─── CV Strategy guidance ─────────────────────────────────────────────────────
with st.expander("💡 ¿Cómo sacar el máximo provecho? Lee esto primero", expanded=False):
    st.markdown("""
**La estrategia correcta: CV Maestro + CVs enfocados**

La mayoría de la gente comete dos errores opuestos:
- **Error 1:** Subir un CV de 10+ páginas con toda su vida — los reclutadores no lo leen y los ATS se pierden.
- **Error 2:** Tener un solo CV genérico que mandan a todas las ofertas — los ATS lo filtran porque no hay match de keywords.

**La solución que esta app facilita:**

1. **Sube tu CV Maestro** — ese documento de 3, 5 o incluso 35 páginas con TODO lo que has hecho. No lo edites, súbelo completo. Claude lo leerá entero y seleccionará solo lo relevante.

2. **Pega la oferta específica** — Claude identifica qué experiencias del maestro encajan, integra las keywords exactas del puesto y descarta lo irrelevante.

3. **Descarga un CV de 1-2 páginas** listo para esa oferta, con formato ATS-friendly.

4. **Repite para cada oferta** — misma subida del maestro, distinta oferta, distinto CV enfocado. Cada uno estará optimizado para esa empresa específica.

**¿Por qué no subir el CV largo directamente a las empresas?**
Los sistemas ATS escanean en segundos. Un CV de 5+ páginas generalmente se descalifica automáticamente antes de que un humano lo vea. El reclutador promedio dedica **7 segundos** al primer escaneo.
    """)

# ─── Inputs ───────────────────────────────────────────────────────────────────
col1, col2 = st.columns(2)

with col1:
    st.subheader("📄 Tu CV")
    st.markdown('<div class="tip-box">💡 <strong>Sube tu CV completo</strong> — aunque tenga 10 o 30 páginas. Claude extraerá solo lo relevante para la oferta.</div>', unsafe_allow_html=True)
    cv_file = st.file_uploader("Sube tu CV", type=["pdf", "docx"],
                                label_visibility="collapsed")
    cv_text_manual = st.text_area("O pega el texto aquí", height=180,
        placeholder="Pega el contenido de tu CV si no tienes archivo...")

with col2:
    st.subheader("💼 Oferta Laboral")
    job_url = st.text_input("🔗 Link de la oferta",
        placeholder="https://www.linkedin.com/jobs/... o cualquier portal")
    job_description = st.text_area("O pega el texto aquí", height=215,
        placeholder="Pega aquí el texto de la oferta...\nMientras más completa, mejor la optimización.")

# ─── Template selector ────────────────────────────────────────────────────────
st.subheader("🎨 Template del CV")
st.markdown("Elige un template. **Después de optimizar, puedes descargar en cualquier template sin costo adicional.**")

TEMPLATES = {
    "📋 Clásico": {
        "desc": "Nombre centrado · Líneas azules · Bullets clásicos",
        "ideal": "Finanzas, legal, gobierno, roles senior",
        "preview": (
            "─────────────────────────────\n"
            "         JUAN PÉREZ          \n"
            "   Gerente de Operaciones    \n"
            "  email | tel | ciudad       \n"
            "─────────────────────────────\n"
            "RESUMEN PROFESIONAL          \n"
            "─────────────────────────────\n"
            "EXPERIENCIA PROFESIONAL      \n"
            "  Cargo · Empresa | 2020-Actual\n"
            "  • Logro con impacto medible\n"
            "─────────────────────────────\n"
            "EDUCACIÓN  /  HABILIDADES    "
        )
    },
    "✨ Moderno": {
        "desc": "Nombre en mayúsculas · Header navy/teal · Flechas ▸",
        "ideal": "Tech, startups, marketing digital, diseño",
        "preview": (
            "JUAN PÉREZ                   \n"
            "Gerente de Operaciones       \n"
            "✉ email  ✆ tel  ⌖ ciudad    \n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "◆  PERFIL PROFESIONAL        \n"
            "◆  EXPERIENCIA               \n"
            "   Cargo — Empresa           \n"
            "   2020-Actual               \n"
            "   ▸ Logro con keywords ATS  \n"
            "◆  HABILIDADES  ◆  IDIOMAS   "
        )
    },
    "🏛️ Ejecutivo": {
        "desc": "Header con franja oscura · Dos columnas de contacto · Serifas",
        "ideal": "Dirección general, C-level, consultoría senior, banca",
        "preview": (
            "████████████████████████████\n"
            "  JUAN PÉREZ                \n"
            "  Director General          \n"
            "████████████████████████████\n"
            "email · tel · LinkedIn      \n"
            "────────────────────────────\n"
            "PERFIL EJECUTIVO            \n"
            "TRAYECTORIA PROFESIONAL     \n"
            "  ■ Cargo  |  Empresa       \n"
            "    Logro estratégico clave \n"
            "────────────────────────────\n"
            "FORMACIÓN  ·  COMPETENCIAS  "
        )
    },
    "⬜ Minimalista": {
        "desc": "Todo en gris oscuro · Sin colores · Máxima legibilidad ATS",
        "ideal": "Cualquier sector · El más seguro para parsers ATS",
        "preview": (
            "Juan Pérez                   \n"
            "Gerente de Operaciones       \n"
            "email | tel | ciudad         \n"
            "                             \n"
            "RESUMEN                      \n"
            "                             \n"
            "EXPERIENCIA                  \n"
            "Cargo — Empresa (2020-Actual)\n"
            "- Logro cuantificado         \n"
            "                             \n"
            "EDUCACIÓN · HABILIDADES      "
        )
    }
}

if "template_choice" not in st.session_state:
    st.session_state["template_choice"] = list(TEMPLATES.keys())[0]

t_cols = st.columns(4)
for i, (name, info) in enumerate(TEMPLATES.items()):
    is_sel = st.session_state["template_choice"] == name
    border_col = "#1B6CA8" if is_sel else "#CCCCCC"
    bg_col = "#E8F4FD" if is_sel else "#FAFAFA"
    check = "✅ " if is_sel else ""
    preview_html = info["preview"].replace("\n", "<br>")
    with t_cols[i]:
        st.markdown(
            f"""<div style="border:2px solid {border_col};border-radius:10px;
                padding:0.8rem;background:{bg_col};min-height:260px;">
            <div style="font-weight:700;font-size:0.95rem;margin-bottom:0.2rem;">
                {check}{name}</div>
            <div style="font-size:0.75rem;color:#888;margin-bottom:0.5rem;">
                {info["ideal"]}</div>
            <div style="font-family:monospace;font-size:0.68rem;color:#555;
                background:#F0F0F0;border-radius:4px;padding:0.4rem 0.5rem;
                line-height:1.5;">{preview_html}</div>
            </div>""",
            unsafe_allow_html=True
        )
        st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)
        if st.button(f"{'✅ Seleccionado' if is_sel else 'Seleccionar'}", key=f"tpl_{i}",
                     use_container_width=True,
                     type="primary" if is_sel else "secondary"):
            st.session_state["template_choice"] = name
            st.rerun()

template = st.session_state.get("template_choice", list(TEMPLATES.keys())[0])

st.markdown("---")

# ─── Helpers ──────────────────────────────────────────────────────────────────
def is_valid_url(url: str) -> bool:
    return bool(re.match(r'^https?://', url.strip()))

def scrape_job_url(url: str) -> str:
    headers = {"User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    )}
    try:
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
    except requests.exceptions.Timeout:
        raise ValueError("El sitio tardó demasiado. Pega el texto manualmente.")
    except requests.exceptions.HTTPError as e:
        raise ValueError(f"No se pudo acceder ({e.response.status_code}).")
    except Exception:
        raise ValueError("No se pudo acceder al link.")

    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script","style","nav","footer","header","aside","form","noscript","iframe"]):
        tag.decompose()
    text = ""
    for sel in [
        {"class": lambda c: c and any(k in " ".join(c).lower() for k in
            ["job-description","description","vacancy","posting","details","content"])},
        {"id": lambda i: i and any(k in i.lower() for k in
            ["job-description","description","details"])},
    ]:
        block = soup.find(attrs=sel)
        if block:
            text = block.get_text(separator="\n", strip=True)
            if len(text) > 200:
                break
    if not text or len(text) < 200:
        text = soup.get_text(separator="\n", strip=True)
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    return "\n".join(lines)[:8000]

def extract_pdf(file) -> str:
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    return text.strip()

def extract_docx(file) -> str:
    doc = DocxDocument(file)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

# ─── Claude optimization ──────────────────────────────────────────────────────
def optimize_cv(api_key, cv_text, job_text, max_pages, font_size) -> dict:
    was_truncated = len(cv_text) > MAX_CV_CHARS
    cv_text = cv_text[:MAX_CV_CHARS]

    words_per_page = {9: 700, 10: 600, 10.5: 560, 11: 520, 12: 460}
    max_words = words_per_page.get(float(font_size), 580) * max_pages

    prompt = f"""Eres un experto coach de carrera y especialista en optimización de CVs para sistemas ATS.

INSTRUCCIONES CRÍTICAS:
1. El CV puede ser un "CV Maestro" extenso con toda la carrera del candidato — selecciona SOLO la experiencia más relevante para ESTA oferta
2. NO inventes nada. Solo reescribe y reorganiza lo que existe
3. Integra las keywords exactas de la oferta de forma natural en logros y resumen
4. Cuantifica logros cuando el CV original tenga datos (números, %, equipos, presupuestos)
5. Respeta ESTRICTAMENTE el límite: {max_pages} página(s) ≈ {max_words} palabras totales
6. El resultado debe pasar parsers ATS: sin tablas, sin columnas, sin elementos gráficos

CV MAESTRO (puede ser extenso — selecciona lo relevante):
{cv_text}

OFERTA DE TRABAJO:
{job_text}

Responde ÚNICAMENTE con JSON válido, sin backticks ni texto adicional:
{{
  "nombre": "nombre completo del candidato",
  "titulo_profesional": "título exacto del puesto ofrecido",
  "email": "email o vacío",
  "telefono": "teléfono o vacío",
  "linkedin": "URL LinkedIn o vacío",
  "ubicacion": "ciudad, país",
  "resumen_profesional": "3-4 oraciones conectando experiencia con requisitos clave. Keywords ATS incluidas naturalmente.",
  "experiencia": [
    {{
      "empresa": "nombre empresa",
      "cargo": "cargo",
      "periodo": "mes/año - mes/año o Actual",
      "logros": [
        "Verbo acción + resultado cuantificado + keyword ATS de la oferta",
        "Solo incluir cargos y logros relevantes para ESTA oferta",
        "Máximo 3-4 logros por cargo"
      ]
    }}
  ],
  "educacion": [
    {{"institucion": "nombre", "titulo": "título", "periodo": "años", "detalle": "mención o vacío"}}
  ],
  "habilidades_tecnicas": ["solo skills del CV que aparecen o son directamente relevantes para la oferta"],
  "habilidades_blandas": ["máx 4 competencias que conectan con la oferta"],
  "idiomas": ["Español - Nativo"],
  "certificaciones": ["solo si existen en el CV original"],
  "ats_compatible": true,
  "ats_razon": "Una frase: por qué pasa o no pasa filtros ATS automáticos",
  "score_match": 82,
  "score_desglose": {{"keywords": 88, "experiencia": 80, "educacion": 75, "habilidades": 85}},
  "score_explicacion": "2-3 oraciones: qué fortalece la candidatura para este rol y qué la limita.",
  "keywords_integradas": ["keyword de la oferta que se integró exitosamente"],
  "keywords_faltantes": ["keyword importante de la oferta ausente en el CV"],
  "coaching": [
    {{"categoria": "Tu fortaleza clave 💪", "tip": "Qué tiene el candidato que es genuinamente valioso para este rol y cómo maximizarlo."}},
    {{"categoria": "Brecha crítica 🎯", "tip": "La brecha más importante. Cómo cerrarla: nombre de curso específico, certificación concreta o proyecto puntual."}},
    {{"categoria": "Quick win de hoy ⚡", "tip": "Una acción en menos de 1 hora que mejore la candidatura para este puesto específico."}},
    {{"categoria": "LinkedIn / Marca personal 🔗", "tip": "Qué cambiar o agregar en LinkedIn para este tipo de rol y empresa."}},
    {{"categoria": "Antes de la entrevista 📋", "tip": "Qué investigar de la empresa, qué preguntas preparar y qué narrativa construir para este puesto."}}
  ]
}}"""

    client = anthropic.Anthropic(api_key=api_key)

    def call_model(model_id: str) -> dict:
        for attempt in range(2):
            msg = client.messages.create(
                model=model_id,
                max_tokens=5000,
                messages=[{"role": "user", "content": prompt}]
            )
            raw = msg.content[0].text.strip()
            if "```json" in raw:
                raw = raw.split("```json")[1].split("```")[0].strip()
            elif "```" in raw:
                raw = raw.split("```")[1].split("```")[0].strip()
            try:
                return json.loads(raw)
            except json.JSONDecodeError:
                if attempt == 0:
                    time.sleep(1)
                    continue
                raise

    # Haiku por defecto (~$0.01). Opus solo si score < 60 (~$0.15)
    result = call_model("claude-haiku-4-5-20251001")
    result["_was_truncated"] = was_truncated
    result["_model_used"] = "haiku"

    if result.get("score_match", 100) < 60:
        result = call_model("claude-opus-4-5-20251101")
        result["_was_truncated"] = was_truncated
        result["_model_used"] = "opus"

    return result

# ─── DOCX helpers ─────────────────────────────────────────────────────────────
def add_border(doc, body_size, border_color):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(body_size)
    p.paragraph_format.space_after = Pt(body_size * 0.3)
    return p

def section_border(p, border_color):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.append(pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), border_color)
    pBdr.append(bottom)

def add_section_header(doc, title, color_rgb, fn, body_size, border_color, prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(body_size)
    p.paragraph_format.space_after = Pt(body_size * 0.3)
    run = p.add_run(f"{prefix}{title}")
    run.bold = True
    run.font.name = fn
    run.font.size = Pt(body_size + 1)
    run.font.color.rgb = RGBColor(*color_rgb)
    section_border(p, border_color)

def R(run, fn, fs, bold=False, italic=False, color=None):
    run.font.name = fn
    run.font.size = Pt(float(fs))
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

# ─── Template: Clásico ────────────────────────────────────────────────────────
def build_clasico(cv, fn, fs):
    fs = float(fs)
    doc = DocxDocument()
    s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(1.0)
    s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)
    DARK=(0x1A,0x1A,0x2E); BLUE=(0x2E,0x75,0xB6); GRAY=(0x66,0x66,0x66)

    def hdr(t): add_section_header(doc, t, BLUE, fn, fs, "2E75B6")

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    R(p.add_run(cv.get("nombre","")), fn, fs+9, bold=True, color=DARK)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    R(p.add_run(cv.get("titulo_profesional","")), fn, fs+2, bold=True, color=BLUE)
    parts = [x for x in [cv.get("email"),cv.get("telefono"),cv.get("ubicacion"),cv.get("linkedin")] if x]
    if parts:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(p.add_run("  |  ".join(parts)), fn, fs-1, color=GRAY)

    if cv.get("resumen_profesional"):
        hdr("RESUMEN PROFESIONAL")
        R(doc.add_paragraph().add_run(cv["resumen_profesional"]), fn, fs)

    if cv.get("experiencia"):
        hdr("EXPERIENCIA PROFESIONAL")
        for exp in cv["experiencia"]:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(fs*0.6)
            R(p.add_run(exp.get("cargo","")), fn, fs, bold=True)
            p2 = doc.add_paragraph()
            R(p2.add_run(f"{exp.get('empresa','')}   |   {exp.get('periodo','')}"), fn, fs-1, italic=True, color=GRAY)
            for logro in exp.get("logros",[]):
                pb = doc.add_paragraph(style="List Bullet")
                pb.paragraph_format.left_indent = Inches(0.2)
                pb.paragraph_format.space_after = Pt(2)
                R(pb.add_run(logro), fn, fs)

    if cv.get("educacion"):
        hdr("EDUCACIÓN")
        for edu in cv["educacion"]:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(fs*0.5)
            R(p.add_run(edu.get("titulo","")), fn, fs, bold=True)
            p2 = doc.add_paragraph()
            R(p2.add_run(f"{edu.get('institucion','')}   |   {edu.get('periodo','')}"), fn, fs-1, italic=True, color=GRAY)
            if edu.get("detalle"): R(doc.add_paragraph().add_run(edu["detalle"]), fn, fs-1)

    if cv.get("habilidades_tecnicas"):
        hdr("HABILIDADES TÉCNICAS")
        R(doc.add_paragraph().add_run("  •  ".join(cv["habilidades_tecnicas"])), fn, fs)
    if cv.get("habilidades_blandas"):
        hdr("COMPETENCIAS")
        R(doc.add_paragraph().add_run("  •  ".join(cv["habilidades_blandas"])), fn, fs)
    if cv.get("idiomas"):
        hdr("IDIOMAS")
        R(doc.add_paragraph().add_run("  |  ".join(cv["idiomas"])), fn, fs)
    certs = [c for c in cv.get("certificaciones",[]) if c]
    if certs:
        hdr("CERTIFICACIONES")
        for cert in certs: R(doc.add_paragraph().add_run(f"• {cert}"), fn, fs)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf

# ─── Template: Moderno ────────────────────────────────────────────────────────
def build_moderno(cv, fn, fs):
    fs = float(fs)
    doc = DocxDocument()
    s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(0.75)
    s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.8)
    NAVY=(0x1B,0x4F,0x72); TEAL=(0x17,0x8A,0xCA); GRAY=(0x77,0x77,0x77)

    def hdr(t): add_section_header(doc, t, NAVY, fn, fs, "17A8CA", prefix="◆  ")

    R(doc.add_paragraph().add_run(cv.get("nombre","").upper()), fn, fs+11, bold=True, color=NAVY)
    R(doc.add_paragraph().add_run(cv.get("titulo_profesional","")), fn, fs+2, bold=True, color=TEAL)
    parts = []
    for icon,key in [("✉","email"),("✆","telefono"),("⌖","ubicacion"),("in","linkedin")]:
        if cv.get(key): parts.append(f"{icon} {cv[key]}")
    if parts: R(doc.add_paragraph().add_run("   |   ".join(parts)), fn, fs-1, color=GRAY)

    p_div = doc.add_paragraph(); p_div.paragraph_format.space_after = Pt(8)
    pPr = p_div._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr'); pPr.append(pBdr)
    btm = OxmlElement('w:bottom')
    for a,v in [('w:val','single'),('w:sz','16'),('w:space','1'),('w:color','1B4F72')]: btm.set(qn(a),v)
    pBdr.append(btm)

    if cv.get("resumen_profesional"):
        hdr("PERFIL PROFESIONAL")
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.15)
        R(p.add_run(cv["resumen_profesional"]), fn, fs)
    if cv.get("experiencia"):
        hdr("EXPERIENCIA")
        for exp in cv["experiencia"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(fs*0.7); p.paragraph_format.left_indent = Inches(0.15)
            R(p.add_run(exp.get("cargo","")), fn, fs, bold=True, color=NAVY)
            R(p.add_run("  —  "), fn, fs)
            R(p.add_run(exp.get("empresa","")), fn, fs)
            p2 = doc.add_paragraph(); p2.paragraph_format.left_indent = Inches(0.15)
            R(p2.add_run(exp.get("periodo","")), fn, fs-1, italic=True, color=TEAL)
            for logro in exp.get("logros",[]):
                pb = doc.add_paragraph(); pb.paragraph_format.left_indent = Inches(0.35); pb.paragraph_format.space_after = Pt(2)
                R(pb.add_run(f"▸  {logro}"), fn, fs)
    if cv.get("educacion"):
        hdr("EDUCACIÓN")
        for edu in cv["educacion"]:
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.15); p.paragraph_format.space_before = Pt(fs*0.5)
            R(p.add_run(edu.get("titulo","")), fn, fs, bold=True, color=NAVY)
            R(p.add_run(f"   |   {edu.get('institucion','')}   |   {edu.get('periodo','')}"), fn, fs-1)
            if edu.get("detalle"):
                p2 = doc.add_paragraph(); p2.paragraph_format.left_indent = Inches(0.15)
                R(p2.add_run(edu["detalle"]), fn, fs-1)
    if cv.get("habilidades_tecnicas") or cv.get("habilidades_blandas"):
        hdr("HABILIDADES")
        for label,key in [("Técnicas: ","habilidades_tecnicas"),("Competencias: ","habilidades_blandas")]:
            if cv.get(key):
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.15)
                R(p.add_run(label), fn, fs, bold=True)
                R(p.add_run("  •  ".join(cv[key])), fn, fs)
    if cv.get("idiomas"):
        hdr("IDIOMAS")
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.15)
        R(p.add_run("  |  ".join(cv["idiomas"])), fn, fs)
    certs = [c for c in cv.get("certificaciones",[]) if c]
    if certs:
        hdr("CERTIFICACIONES")
        for cert in certs:
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.15)
            R(p.add_run(f"▸  {cert}"), fn, fs)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf

# ─── Template: Ejecutivo ──────────────────────────────────────────────────────
def build_ejecutivo(cv, fn, fs):
    fs = float(fs)
    # Ejecutivo uses Georgia for body if Calibri/Arial selected (more gravitas)
    serif_fn = "Georgia" if fn in ["Calibri","Arial","Trebuchet MS"] else fn
    doc = DocxDocument()
    s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(1.0)
    s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.8)
    NAVY=(0x1B,0x2A,0x4A); GOLD=(0x8B,0x6C,0x1E); DARK=(0x22,0x22,0x22); GRAY=(0x55,0x55,0x55)

    # Name block — bold, large, left aligned
    p = doc.add_paragraph()
    R(p.add_run(cv.get("nombre","").upper()), serif_fn, fs+10, bold=True, color=NAVY)
    p2 = doc.add_paragraph()
    R(p2.add_run(cv.get("titulo_profesional","")), serif_fn, fs+1, italic=True, color=GOLD)

    # Contact line
    parts = [x for x in [cv.get("email"),cv.get("telefono"),cv.get("ubicacion"),cv.get("linkedin")] if x]
    if parts:
        p3 = doc.add_paragraph()
        R(p3.add_run("  ·  ".join(parts)), fn, fs-1, color=GRAY)

    # Thick navy divider
    p_div = doc.add_paragraph(); p_div.paragraph_format.space_after = Pt(6)
    pPr = p_div._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr'); pPr.append(pBdr)
    btm = OxmlElement('w:bottom')
    for a,v in [('w:val','single'),('w:sz','24'),('w:space','1'),('w:color','1B2A4A')]: btm.set(qn(a),v)
    pBdr.append(btm)

    def hdr(t): add_section_header(doc, t, NAVY, serif_fn, fs, "1B2A4A")

    if cv.get("resumen_profesional"):
        hdr("PERFIL EJECUTIVO")
        p = doc.add_paragraph()
        R(p.add_run(cv["resumen_profesional"]), serif_fn, fs, italic=True)

    if cv.get("experiencia"):
        hdr("TRAYECTORIA PROFESIONAL")
        for exp in cv["experiencia"]:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(fs*0.7)
            R(p.add_run(f"■  {exp.get('cargo','')}"), serif_fn, fs, bold=True, color=NAVY)
            R(p.add_run(f"  ·  {exp.get('empresa','')}"), serif_fn, fs, color=DARK)
            p2 = doc.add_paragraph()
            R(p2.add_run(exp.get("periodo","")), fn, fs-1, italic=True, color=GRAY)
            p2.paragraph_format.left_indent = Inches(0.25)
            for logro in exp.get("logros",[]):
                pb = doc.add_paragraph()
                pb.paragraph_format.left_indent = Inches(0.35); pb.paragraph_format.space_after = Pt(2)
                R(pb.add_run(f"›  {logro}"), serif_fn, fs)

    if cv.get("educacion"):
        hdr("FORMACIÓN ACADÉMICA")
        for edu in cv["educacion"]:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(fs*0.5)
            R(p.add_run(edu.get("titulo","")), serif_fn, fs, bold=True, color=NAVY)
            R(p.add_run(f"  —  {edu.get('institucion','')}  |  {edu.get('periodo','')}"), serif_fn, fs-1, color=GRAY)
            if edu.get("detalle"): R(doc.add_paragraph().add_run(edu["detalle"]), serif_fn, fs-1, italic=True)

    # Skills in two-column feel using tab stops
    if cv.get("habilidades_tecnicas") or cv.get("habilidades_blandas"):
        hdr("COMPETENCIAS Y HABILIDADES")
        all_skills = cv.get("habilidades_tecnicas",[]) + cv.get("habilidades_blandas",[])
        p = doc.add_paragraph()
        R(p.add_run("  ■  ".join(all_skills)), serif_fn, fs)

    if cv.get("idiomas"):
        hdr("IDIOMAS")
        R(doc.add_paragraph().add_run("  |  ".join(cv["idiomas"])), serif_fn, fs)

    certs = [c for c in cv.get("certificaciones",[]) if c]
    if certs:
        hdr("CERTIFICACIONES")
        for cert in certs: R(doc.add_paragraph().add_run(f"■  {cert}"), serif_fn, fs)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf

# ─── Template: Minimalista ────────────────────────────────────────────────────
def build_minimalista(cv, fn, fs):
    fs = float(fs)
    doc = DocxDocument()
    s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(1.1)
    s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
    DARK=(0x22,0x22,0x22); MID=(0x44,0x44,0x44); LIGHT=(0x88,0x88,0x88)

    # Name — simple, no color
    p = doc.add_paragraph()
    R(p.add_run(cv.get("nombre","")), fn, fs+6, bold=True, color=DARK)
    p2 = doc.add_paragraph()
    R(p2.add_run(cv.get("titulo_profesional","")), fn, fs+0.5, color=MID)

    parts = [x for x in [cv.get("email"),cv.get("telefono"),cv.get("ubicacion"),cv.get("linkedin")] if x]
    if parts:
        p3 = doc.add_paragraph()
        R(p3.add_run("  |  ".join(parts)), fn, fs-1, color=LIGHT)

    def hdr(t):
        # Simple uppercase, no color, thin border
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(fs*1.2)
        p.paragraph_format.space_after = Pt(fs*0.2)
        run = p.add_run(t.upper())
        run.bold = True; run.font.name = fn; run.font.size = Pt(fs)
        run.font.color.rgb = RGBColor(*DARK)
        section_border(p, "AAAAAA")

    if cv.get("resumen_profesional"):
        hdr("Resumen")
        R(doc.add_paragraph().add_run(cv["resumen_profesional"]), fn, fs, color=MID)

    if cv.get("experiencia"):
        hdr("Experiencia")
        for exp in cv["experiencia"]:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(fs*0.5)
            R(p.add_run(exp.get("cargo","")), fn, fs, bold=True, color=DARK)
            R(p.add_run(f"  —  {exp.get('empresa','')}"), fn, fs, color=MID)
            p2 = doc.add_paragraph()
            R(p2.add_run(exp.get("periodo","")), fn, fs-1, italic=True, color=LIGHT)
            for logro in exp.get("logros",[]):
                pb = doc.add_paragraph()
                pb.paragraph_format.left_indent = Inches(0.2); pb.paragraph_format.space_after = Pt(2)
                R(pb.add_run(f"- {logro}"), fn, fs, color=MID)

    if cv.get("educacion"):
        hdr("Educación")
        for edu in cv["educacion"]:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(fs*0.4)
            R(p.add_run(edu.get("titulo","")), fn, fs, bold=True, color=DARK)
            R(p.add_run(f"  —  {edu.get('institucion','')}  ({edu.get('periodo','')})"), fn, fs-1, color=MID)
            if edu.get("detalle"): R(doc.add_paragraph().add_run(edu["detalle"]), fn, fs-1, color=LIGHT)

    all_skills = cv.get("habilidades_tecnicas",[]) + cv.get("habilidades_blandas",[])
    if all_skills:
        hdr("Habilidades")
        R(doc.add_paragraph().add_run("  /  ".join(all_skills)), fn, fs, color=MID)

    if cv.get("idiomas"):
        hdr("Idiomas")
        R(doc.add_paragraph().add_run("  |  ".join(cv["idiomas"])), fn, fs, color=MID)

    certs = [c for c in cv.get("certificaciones",[]) if c]
    if certs:
        hdr("Certificaciones")
        R(doc.add_paragraph().add_run("  /  ".join(certs)), fn, fs, color=MID)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf

# ─── Template dispatcher ──────────────────────────────────────────────────────
BUILDERS = {
    "📋 Clásico":    build_clasico,
    "✨ Moderno":    build_moderno,
    "🏛️ Ejecutivo":  build_ejecutivo,
    "⬜ Minimalista": build_minimalista,
}

# ─── Results display ──────────────────────────────────────────────────────────
def show_results(cv_data, template, fn, fs, max_pages):
    st.markdown("---")
    st.subheader("📊 Análisis de Compatibilidad")

    if cv_data.get("_was_truncated"):
        st.markdown(
            '<div class="warn-box">⚠️ Tu CV era muy extenso — se analizaron los primeros 15.000 caracteres, '
            'que equivalen a unas 5-6 páginas densas. Si tienes experiencia muy antigua que no apareció, '
            'probablemente es porque Claude la priorizó correctamente como no relevante para esta oferta.</div>',
            unsafe_allow_html=True
        )

    ats_ok  = cv_data.get("ats_compatible", True)
    ats_msg = cv_data.get("ats_razon", "")
    score   = cv_data.get("score_match", 0)
    sc_col  = "🟢" if score >= 75 else "🟡" if score >= 55 else "🔴"

    bc, sc = st.columns([1, 2])
    with bc:
        if ats_ok: st.success("✅ ATS Compatible")
        else: st.error("❌ No ATS Compatible")
        if ats_msg: st.caption(ats_msg)
    with sc:
        st.metric(f"{sc_col} Match con la oferta (estimado por IA)", f"{score}%")
        explain = cv_data.get("score_explicacion","")
        if explain:
            st.markdown(f'<div class="score-explain">{explain}</div>', unsafe_allow_html=True)

    desglose = cv_data.get("score_desglose",{})
    if desglose:
        with st.expander("📈 Ver desglose del score"):
            d1,d2,d3,d4 = st.columns(4)
            d1.metric("Keywords",    f"{desglose.get('keywords','-')}%")
            d2.metric("Experiencia", f"{desglose.get('experiencia','-')}%")
            d3.metric("Educación",   f"{desglose.get('educacion','-')}%")
            d4.metric("Habilidades", f"{desglose.get('habilidades','-')}%")

    st.markdown("---")
    k1,k2 = st.columns(2)
    with k1:
        kw_ok = cv_data.get("keywords_integradas",[])
        if kw_ok: st.success(f"✅ **Keywords integradas ({len(kw_ok)}):**\n" + ", ".join(kw_ok))
    with k2:
        kw_miss = cv_data.get("keywords_faltantes",[])
        if kw_miss: st.warning(f"⚠️ **Keywords ausentes ({len(kw_miss)}):**\n" + ", ".join(kw_miss))

    coaching = cv_data.get("coaching",[])
    if coaching:
        st.markdown("---")
        st.subheader("🎯 Tu Plan de Acción")
        st.markdown("Recomendaciones personalizadas para maximizar tus chances en **esta** postulación:")
        for tip_item in coaching:
            cat = tip_item.get("categoria","")
            tip_txt = tip_item.get("tip","")
            st.markdown(f'<div class="coach-card"><strong>{cat}</strong><br>{tip_txt}</div>',
                        unsafe_allow_html=True)

    # Download section — all 4 templates
    st.markdown("---")
    st.subheader("⬇️ Descarga tu CV")
    st.markdown("**El análisis ya está hecho — descarga en cualquier template sin costo adicional:**")

    nombre = cv_data.get("nombre","cv").replace(" ","_")
    dl_cols = st.columns(4)
    for i, (tname, builder) in enumerate(BUILDERS.items()):
        with dl_cols[i]:
            try:
                buf = builder(cv_data, fn, fs)
                short = tname.split(" ")[1]  # "Clásico", "Moderno", etc.
                st.download_button(
                    label=f"⬇️ {short}",
                    data=buf,
                    file_name=f"CV_ATS_{nombre}_{short}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"Error {tname}: {e}")

    st.caption(f"Tipografía: {fn} · {fs}pt · {max_pages} página(s) máxima(s)")

# ─── Regenerate only ──────────────────────────────────────────────────────────
if st.session_state.get("regen_docx") and st.session_state.get("cv_data"):
    st.session_state["regen_docx"] = False
    show_results(st.session_state["cv_data"], template, font_family, font_size, max_pages)
    st.stop()

# ─── Main action ──────────────────────────────────────────────────────────────
if st.button("🚀 Optimizar mi CV", use_container_width=True):
    if not api_key:
        st.error("⚠️ No hay API Key disponible.")
        st.stop()

    # Resolve job text
    final_job = job_description.strip()
    if job_url.strip():
        if not is_valid_url(job_url):
            st.warning("⚠️ El link debe empezar con http:// o https://")
        else:
            with st.spinner("🔍 Leyendo la oferta desde el link..."):
                try:
                    scraped = scrape_job_url(job_url.strip())
                    if scraped:
                        final_job = scraped
                        st.success(f"✅ Oferta leída ({len(scraped):,} caracteres extraídos)")
                    else:
                        st.warning("No se pudo extraer texto del link. Usando el texto pegado.")
                except ValueError as e:
                    st.warning(str(e))

    if not final_job:
        st.error("⚠️ Pega la oferta de trabajo o ingresa un link válido.")
        st.stop()

    # Extract CV
    cv_text = ""
    if cv_file:
        with st.spinner("📄 Extrayendo texto del archivo..."):
            try:
                cv_text = extract_pdf(cv_file) if cv_file.name.lower().endswith(".pdf") \
                          else extract_docx(cv_file)
            except Exception as e:
                st.error(f"Error al leer el archivo: {e}"); st.stop()

    if cv_text_manual.strip():
        cv_text = cv_text_manual.strip() if not cv_text else cv_text + "\n" + cv_text_manual.strip()

    if not cv_text:
        st.error("⚠️ Sube un CV o pega el texto manualmente."); st.stop()

    # Call Claude
    with st.spinner("🤖 Analizando tu CV maestro, keywords de la oferta y preparando coaching... (puede tomar 20-40 segundos)"):
        try:
            cv_data = optimize_cv(api_key, cv_text, final_job, max_pages, font_size)
            st.session_state["cv_data"] = cv_data
            st.session_state["api_credits_error"] = False
        except json.JSONDecodeError:
            st.error("Error procesando respuesta de Claude. Intenta nuevamente."); st.stop()
        except anthropic.AuthenticationError:
            st.error("API Key inválida."); st.stop()
        except anthropic.RateLimitError:
            st.session_state["api_credits_error"] = True
            st.error("⚠️ Servicio sin saldo. Ingresa tu API Key en el panel lateral.")
            st.rerun()
        except Exception as e:
            st.error(f"Error inesperado: {e}"); st.stop()

    show_results(cv_data, template, font_family, font_size, max_pages)

st.markdown("---")
st.caption("CV Optimizer ATS · Powered by Claude AI · Anthropic")
