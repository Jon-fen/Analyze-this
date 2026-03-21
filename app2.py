# CV Optimizer ATS — app2.py (v2, con usuarios y créditos)
# Requiere en Streamlit Secrets:
#   ANTHROPIC_API_KEY, SUPABASE_URL, SUPABASE_KEY

import streamlit as st
import anthropic
import pdfplumber
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors as rl_colors
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.platypus import Image as RLImage
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io, json, re, time
import requests
from bs4 import BeautifulSoup
from supabase import create_client, Client
from datetime import datetime, timezone
from streamlit_oauth import OAuth2Component

MAX_CV_CHARS = 40_000   # ~15 páginas densas
# Social proof base offset — real usage + base for credibility
COUNTER_BASE_CVS   = 10_000
COUNTER_BASE_USERS = 1_000
MAX_CV_CHARS_CAREER = 80_000  # Sin límite para cambio de carrera

# ─── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(page_title="CV Optimizer ATS", page_icon="🎯", layout="centered", initial_sidebar_state="expanded")

st.markdown("""
<style>
  @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

  /* ── Tokens ── */
  :root {
    --bg:        #0a0a0b;
    --bg2:       #111114;
    --bg3:       #18181c;
    --border:    rgba(255,255,255,0.07);
    --border2:   rgba(255,255,255,0.12);
    --text:      #e4e4e7;
    --text2:     #a1a1aa;
    --text3:     #71717a;
    --blue:      #3b82f6;
    --blue-dim:  rgba(59,130,246,0.12);
    --blue-bdr:  rgba(59,130,246,0.25);
    --green:     #22c55e;
    --green-dim: rgba(34,197,94,0.10);
    --amber:     #f59e0b;
    --amber-dim: rgba(245,158,11,0.10);
    --red:       #ef4444;
    --radius:    8px;
    --radius-lg: 12px;
    color-scheme: dark;
  }

  /* ── Base ── */
  html, body, .stApp, [data-testid="stAppViewContainer"],
  [data-testid="stMain"], [data-testid="block-container"] {
    background-color: var(--bg) !important;
    color: var(--text) !important;
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif !important;
  }

  /* ── Layout ── */
  .block-container {
    padding-top: 1.5rem !important;
    padding-bottom: 3rem !important;
    max-width: 860px !important;
  }

  /* ── Sidebar ── */
  section[data-testid="stSidebar"] {
    background-color: var(--bg2) !important;
    border-right: 1px solid var(--border) !important;
  }
  section[data-testid="stSidebar"] > div { padding-top: 1rem !important; }

  /* ── Hide Streamlit chrome ── */
  #MainMenu, footer, header, [data-testid="stToolbar"] { display: none !important; }

  /* ── Typography ── */
  h1, h2, h3 { font-family: 'Inter', sans-serif !important; letter-spacing: -0.025em !important; }

  /* ── Buttons — primary ── */
  .stButton > button[kind="primary"],
  button[data-testid="baseButton-primary"] {
    background: var(--blue) !important;
    color: #fff !important;
    border: none !important;
    border-radius: var(--radius) !important;
    font-weight: 600 !important;
    font-size: 0.9rem !important;
    padding: 0.55rem 1.2rem !important;
    transition: background 0.15s, transform 0.1s !important;
    box-shadow: 0 0 0 0 transparent !important;
  }
  .stButton > button[kind="primary"]:hover { background: #2563eb !important; }
  .stButton > button[kind="primary"]:active { transform: scale(0.98) !important; }

  /* ── Buttons — secondary (default) ── */
  .stButton > button,
  button[data-testid="baseButton-secondary"] {
    background: var(--bg3) !important;
    color: var(--text) !important;
    border: 1px solid var(--border2) !important;
    border-radius: var(--radius) !important;
    font-weight: 500 !important;
    font-size: 0.875rem !important;
    padding: 0.5rem 1rem !important;
    transition: background 0.15s, border-color 0.15s !important;
  }
  .stButton > button:hover { background: #222226 !important; border-color: rgba(255,255,255,0.2) !important; }

  /* ── Download buttons ── */
  div[data-testid="stDownloadButton"] button {
    background: var(--bg3) !important;
    color: var(--blue) !important;
    border: 1px solid var(--blue-bdr) !important;
    border-radius: var(--radius) !important;
    font-weight: 600 !important;
    font-size: 0.85rem !important;
    width: 100% !important;
    padding: 0.5rem 0.8rem !important;
    transition: background 0.15s !important;
  }
  div[data-testid="stDownloadButton"] button:hover {
    background: var(--blue-dim) !important;
  }

  /* ── Inputs ── */
  .stTextInput input, .stTextArea textarea, .stSelectbox select {
    background: var(--bg2) !important;
    border: 1px solid var(--border2) !important;
    border-radius: var(--radius) !important;
    color: var(--text) !important;
    font-family: 'Inter', sans-serif !important;
    font-size: 0.9rem !important;
    transition: border-color 0.15s !important;
  }
  .stTextInput input:focus, .stTextArea textarea:focus {
    border-color: var(--blue) !important;
    box-shadow: 0 0 0 3px rgba(59,130,246,0.15) !important;
    outline: none !important;
  }

  /* ── File uploader ── */
  [data-testid="stFileUploader"] {
    background: var(--bg2) !important;
    border: 1.5px dashed var(--border2) !important;
    border-radius: var(--radius-lg) !important;
    padding: 1rem !important;
    transition: border-color 0.15s !important;
  }
  [data-testid="stFileUploader"]:hover { border-color: var(--blue) !important; }

  /* ── Tabs ── */
  [data-testid="stTabs"] [data-baseweb="tab-list"] {
    background: var(--bg2) !important;
    border-radius: var(--radius) !important;
    padding: 3px !important;
    gap: 2px !important;
    border: 1px solid var(--border) !important;
  }
  [data-testid="stTabs"] [data-baseweb="tab"] {
    background: transparent !important;
    border-radius: 6px !important;
    color: var(--text2) !important;
    font-weight: 500 !important;
    font-size: 0.85rem !important;
    padding: 0.4rem 1rem !important;
    transition: all 0.15s !important;
  }
  [data-testid="stTabs"] [aria-selected="true"] {
    background: var(--bg3) !important;
    color: var(--text) !important;
    box-shadow: 0 1px 3px rgba(0,0,0,0.4) !important;
  }
  [data-testid="stTabContent"] {
    border: 1px solid var(--border) !important;
    border-top: none !important;
    border-radius: 0 0 var(--radius) var(--radius) !important;
    padding: 1.25rem !important;
    background: var(--bg2) !important;
  }

  /* ── Expanders ── */
  [data-testid="stExpander"] {
    background: var(--bg2) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius-lg) !important;
    margin-bottom: 0.5rem !important;
    overflow: hidden !important;
    transition: border-color 0.15s !important;
  }
  [data-testid="stExpander"]:hover { border-color: var(--border2) !important; }
  [data-testid="stExpander"] summary {
    font-weight: 500 !important;
    font-size: 0.9rem !important;
    color: var(--text) !important;
    padding: 0.75rem 1rem !important;
  }
  [data-testid="stExpander"] [data-testid="stExpanderDetails"] {
    padding: 0 1rem 1rem 1rem !important;
    font-size: 0.875rem !important;
    color: var(--text2) !important;
    line-height: 1.65 !important;
  }
  /* Coaching color accents */
  [data-testid="stExpander"]:nth-of-type(1) { border-left: 2px solid var(--green) !important; }
  [data-testid="stExpander"]:nth-of-type(2) { border-left: 2px solid var(--amber) !important; }
  [data-testid="stExpander"]:nth-of-type(3) { border-left: 2px solid var(--blue) !important; }
  [data-testid="stExpander"]:nth-of-type(4) { border-left: 2px solid #a78bfa !important; }
  [data-testid="stExpander"]:nth-of-type(5) { border-left: 2px solid #38bdf8 !important; }

  /* ── Toggles ── */
  [data-testid="stToggle"] {
    background: var(--bg2) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius-lg) !important;
    padding: 0.7rem 1rem !important;
    margin-bottom: 0.4rem !important;
    transition: border-color 0.15s !important;
  }
  [data-testid="stToggle"]:hover { border-color: var(--blue-bdr) !important; }
  [data-testid="stToggle"] label, [data-testid="stToggle"] p {
    font-size: 0.9rem !important;
    font-weight: 500 !important;
    color: var(--text) !important;
  }

  /* ── Metrics ── */
  [data-testid="metric-container"] {
    background: var(--bg2) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius-lg) !important;
    padding: 1rem 1.1rem !important;
  }
  [data-testid="metric-container"] [data-testid="stMetricValue"] {
    font-size: 1.6rem !important;
    font-weight: 700 !important;
    color: var(--text) !important;
  }
  [data-testid="metric-container"] [data-testid="stMetricLabel"] {
    font-size: 0.78rem !important;
    color: var(--text2) !important;
    font-weight: 500 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.05em !important;
  }

  /* ── Alerts ── */
  [data-testid="stAlert"] {
    border-radius: var(--radius) !important;
    border-width: 1px !important;
    font-size: 0.875rem !important;
  }

  /* ── Progress bar ── */
  [data-testid="stProgressBar"] > div > div {
    background: var(--blue) !important;
    border-radius: 99px !important;
  }
  [data-testid="stProgressBar"] > div {
    background: var(--bg3) !important;
    border-radius: 99px !important;
    height: 4px !important;
  }

  /* ── Radio ── */
  [data-testid="stRadio"] label { font-size: 0.875rem !important; }

  /* ── Divider ── */
  hr { border-color: var(--border) !important; margin: 1.25rem 0 !important; }

  /* ── Custom component classes ── */
  .at-card {
    background: var(--bg2);
    border: 1px solid var(--border);
    border-radius: var(--radius-lg);
    padding: 1.1rem 1.2rem;
    margin-bottom: 0.5rem;
  }
  .at-card:hover { border-color: var(--border2); }

  .at-badge {
    display: inline-flex; align-items: center; gap: 6px;
    padding: 0.25rem 0.75rem;
    border-radius: 99px;
    font-size: 0.78rem; font-weight: 600;
    letter-spacing: 0.02em;
  }
  .at-badge-blue  { background: var(--blue-dim); color: var(--blue); border: 1px solid var(--blue-bdr); }
  .at-badge-green { background: var(--green-dim); color: var(--green); border: 1px solid rgba(34,197,94,0.25); }
  .at-badge-amber { background: var(--amber-dim); color: var(--amber); border: 1px solid rgba(245,158,11,0.25); }
  .at-badge-red   { background: rgba(239,68,68,0.10); color: var(--red); border: 1px solid rgba(239,68,68,0.25); }

  .score-explain {
    background: var(--bg3);
    border: 1px solid var(--border);
    border-radius: var(--radius);
    padding: 0.75rem 1rem;
    font-size: 0.875rem;
    color: var(--text2);
    margin-top: 0.5rem;
    line-height: 1.6;
  }

  .warn-truncate {
    background: var(--amber-dim);
    border-left: 2px solid var(--amber);
    padding: 0.5rem 0.8rem;
    border-radius: var(--radius);
    font-size: 0.83rem;
    color: var(--amber);
    margin-bottom: 0.75rem;
  }

  .kw-chip-ok {
    display: inline-block; padding: 2px 10px; margin: 2px 3px;
    background: var(--green-dim); color: var(--green);
    border: 1px solid rgba(34,197,94,0.25); border-radius: 99px;
    font-size: 0.78rem; font-weight: 500;
  }
  .kw-chip-miss {
    display: inline-block; padding: 2px 10px; margin: 2px 3px;
    background: var(--amber-dim); color: var(--amber);
    border: 1px solid rgba(245,158,11,0.25); border-radius: 99px;
    font-size: 0.78rem; font-weight: 500;
  }
  .kw-section { margin-bottom: 0.7rem; }
  .kw-label { font-size: 0.75rem; font-weight: 600; margin-bottom: 0.4rem; display: block;
              text-transform: uppercase; letter-spacing: 0.06em; }
  .kw-label-ok   { color: var(--green); }
  .kw-label-miss { color: var(--amber); }

  /* ── ATS live badge ── */
  .ats-live {
    display: inline-flex; align-items: center; gap: 8px;
    padding: 6px 14px; border-radius: 99px;
    background: var(--green-dim);
    border: 1px solid rgba(34,197,94,0.25);
    font-size: 0.82rem; font-weight: 600; color: var(--green);
  }
  .ats-dot {
    width: 7px; height: 7px; border-radius: 50%;
    background: var(--green); flex-shrink: 0;
    animation: ats-pulse 2s ease-out infinite;
  }
  @keyframes ats-pulse {
    0%,100% { box-shadow: 0 0 0 0 rgba(34,197,94,0.4); }
    50%     { box-shadow: 0 0 0 5px rgba(34,197,94,0); }
  }

  /* ── Sidebar user block ── */
  .sb-user {
    display: flex; align-items: center; gap: 10px;
    padding: 0.75rem 0.9rem;
    background: var(--bg3);
    border: 1px solid var(--border);
    border-radius: var(--radius-lg);
    margin-bottom: 0.9rem;
  }
  .sb-avatar {
    width: 34px; height: 34px; border-radius: 50%;
    background: linear-gradient(135deg, #3b82f6 0%, #8b5cf6 100%);
    display: flex; align-items: center; justify-content: center;
    font-size: 0.9rem; font-weight: 700; color: #fff;
    flex-shrink: 0;
  }
  .sb-name { font-size: 0.85rem; font-weight: 600; color: var(--text); }
  .sb-plan { font-size: 0.72rem; color: var(--text3); margin-top: 1px; }

  /* ── Sidebar plan pill ── */
  .sb-pill {
    display: inline-block; padding: 2px 8px;
    border-radius: 99px; font-size: 0.65rem; font-weight: 700;
    letter-spacing: 0.05em; text-transform: uppercase;
  }
  .sb-pill-free  { background: var(--bg3); color: var(--text2); border: 1px solid var(--border2); }
  .sb-pill-pro   { background: rgba(245,158,11,0.12); color: var(--amber); border: 1px solid rgba(245,158,11,0.25); }
  .sb-pill-admin { background: rgba(168,85,247,0.12); color: #c084fc; border: 1px solid rgba(168,85,247,0.25); }

  /* ── Credits bar custom ── */
  .cred-bar-wrap {
    background: var(--bg3); border-radius: 99px; height: 4px;
    margin: 0.4rem 0 0.8rem 0; overflow: hidden;
  }
  .cred-bar-fill {
    height: 100%; border-radius: 99px;
    background: var(--blue);
    transition: width 0.4s ease;
  }
  .cred-bar-fill.low  { background: var(--amber); }
  .cred-bar-fill.zero { background: var(--red); }

  /* ── App header ── */
  .app-header {
    display: flex; align-items: flex-start; justify-content: space-between;
    padding: 0 0 1.2rem 0;
    border-bottom: 1px solid var(--border);
    margin-bottom: 1.5rem;
  }
  .app-logo-row { display: flex; align-items: center; gap: 10px; }
  .app-title {
    font-size: 1.3rem; font-weight: 800; letter-spacing: -0.03em;
    color: var(--text); margin: 0;
  }
  .app-beta {
    font-size: 0.6rem; font-weight: 700; letter-spacing: 0.08em;
    background: var(--amber-dim); color: var(--amber);
    border: 1px solid rgba(245,158,11,0.3);
    padding: 2px 7px; border-radius: 99px;
    vertical-align: middle; text-transform: uppercase;
  }
  .app-subtitle { font-size: 0.83rem; color: var(--text3); margin: 3px 0 0 0; }

  /* ── Section divider with label ── */
  .section-label {
    font-size: 0.7rem; font-weight: 700; letter-spacing: 0.1em;
    text-transform: uppercase; color: var(--text3);
    margin: 1.2rem 0 0.6rem 0;
  }

  /* ── Guest feature pills ── */
  .feat-grid {
    display: grid; grid-template-columns: repeat(3, 1fr); gap: 0.6rem;
    margin-bottom: 1.2rem;
  }
  .feat-pill {
    background: var(--bg2); border: 1px solid var(--border);
    border-radius: var(--radius-lg); padding: 0.8rem 0.7rem; text-align: center;
    transition: border-color 0.15s;
  }
  .feat-pill:hover { border-color: var(--blue-bdr); }
  .feat-pill-icon { font-size: 1.3rem; }
  .feat-pill-title { font-size: 0.78rem; font-weight: 600; color: var(--text); margin-top: 0.3rem; }
  .feat-pill-sub   { font-size: 0.68rem; color: var(--text3); margin-top: 0.15rem; }

  /* ── Template download cards ── */
  .tpl-header {
    text-align: center; padding: 0.5rem 0 0.3rem 0;
    font-size: 0.75rem; color: var(--text2);
  }
  .tpl-icon { font-size: 1.3rem; display: block; margin-bottom: 0.2rem; }
  .tpl-name { font-weight: 600; color: var(--text); font-size: 0.8rem; }
  .tpl-ideal { font-size: 0.68rem; color: var(--text3); margin-top: 0.1rem; }

  /* ── Prevent transition flash ── */
  .stApp { transition: none !important; }

  /* ── Sidebar collapse ── */
  [data-testid="collapsedControl"] { display: flex !important; }

  /* ── Scroll hint ── */
  @keyframes pulse { 0%,100% { opacity:0.4; } 50% { opacity:1; } }
  .scroll-hint { animation: pulse 2.5s infinite; }
</style>
""", unsafe_allow_html=True)

# ─── Secrets ──────────────────────────────────────────────────────────────────
def get_secret(key: str, default=None):
    try:
        return st.secrets[key]
    except (KeyError, FileNotFoundError):
        return default

ANTHROPIC_KEY     = get_secret("ANTHROPIC_API_KEY")
SUPABASE_URL      = get_secret("SUPABASE_URL")
SUPABASE_KEY      = get_secret("SUPABASE_KEY")
GOOGLE_CLIENT_ID  = get_secret("GOOGLE_CLIENT_ID", "")
GOOGLE_CLIENT_SECRET = get_secret("GOOGLE_CLIENT_SECRET", "")

PLAN_CREDITS = {"free": 10, "pro": 50, "admin": 999999}

# ─── Supabase client ──────────────────────────────────────────────────────────
@st.cache_resource
def get_supabase() -> Client:
    if not SUPABASE_URL or not SUPABASE_KEY:
        return None
    return create_client(SUPABASE_URL, SUPABASE_KEY)

supabase = get_supabase()

# ─── Auth helpers ─────────────────────────────────────────────────────────────
def sign_up(email: str, password: str, display_name: str = "") -> tuple[bool, str]:
    try:
        res = supabase.auth.sign_up({"email": email, "password": password})
        if res.user:
            supabase.table("profiles").insert({
                "id": res.user.id,
                "email": email,
                "display_name": display_name or email.split("@")[0],
                "plan": "free",
                "credits_used_this_month": 0,
                "credits_reset_at": datetime.now(timezone.utc).isoformat()
            }).execute()
            return True, "✅ Cuenta creada. Revisa tu email para confirmar."
        return False, "Error al crear cuenta."
    except Exception as e:
        return False, str(e)

def sign_in(email: str, password: str) -> tuple[bool, str]:
    try:
        res = supabase.auth.sign_in_with_password({"email": email, "password": password})
        if res.user:
            st.session_state["user"] = res.user
            st.session_state["session"] = res.session
            if res.session:
                st.session_state["_access_token"]  = res.session.access_token
                st.session_state["_refresh_token"] = res.session.refresh_token
            return True, ""
        return False, "Credenciales incorrectas."
    except Exception as e:
        return False, "Email o contraseña incorrectos."

def restore_session():
    """Restore session from stored tokens — keeps user logged in across reruns.
    Strategy: try set_session first; if that raises (expired access token)
    fall back to refresh_session using only the refresh token.
    """
    if st.session_state.get("user"):
        return
    access  = st.session_state.get("_access_token", "")
    refresh = st.session_state.get("_refresh_token", "")
    if not refresh:
        return
    # ── Attempt 1: set full session (works when access token still valid) ──
    if access:
        try:
            session = supabase.auth.set_session(access, refresh)
            if session and session.user:
                st.session_state["user"]    = session.user
                st.session_state["session"] = session
                if session.session:
                    st.session_state["_access_token"]  = session.session.access_token
                    st.session_state["_refresh_token"] = session.session.refresh_token
                return
        except Exception:
            pass  # access token may be expired — try refresh below
    # ── Attempt 2: refresh using only the refresh token ───────────────────
    try:
        session = supabase.auth.refresh_session(refresh)
        if session and session.user:
            st.session_state["user"]    = session.user
            st.session_state["session"] = session
            if session.session:
                st.session_state["_access_token"]  = session.session.access_token
                st.session_state["_refresh_token"] = session.session.refresh_token
        else:
            # Both attempts failed — clear stale tokens
            st.session_state.pop("_access_token", None)
            st.session_state.pop("_refresh_token", None)
    except Exception:
        st.session_state.pop("_access_token", None)
        st.session_state.pop("_refresh_token", None)

def send_magic_link(email: str) -> tuple[bool, str]:
    try:
        supabase.auth.sign_in_with_otp({
            "email": email,
            "options": {"should_create_user": True,
                        "email_redirect_to": "https://analyze-this-v2.streamlit.app"}
        })
        return True, "✅ Link enviado. Revisa tu email y haz clic para entrar."
    except Exception as e:
        return False, f"Error: {e}"

def handle_magic_callback():
    """Intercept Supabase magic link token from URL — one-time use."""
    if st.session_state.get("_magic_handled"):
        return
    params = dict(st.query_params)
    # Supabase magic link returns access_token + refresh_token as query params
    access_token  = params.get("access_token", "")
    refresh_token = params.get("refresh_token", "")
    if not access_token:
        return
    try:
        session = supabase.auth.set_session(access_token, refresh_token)
        if session and session.user:
            st.session_state["user"]    = session.user
            st.session_state["session"] = session
            st.session_state["_magic_handled"] = True
            if session.session:
                st.session_state["_access_token"]  = session.session.access_token
                st.session_state["_refresh_token"] = session.session.refresh_token
            # Create profile if first time
            try:
                existing = get_profile(session.user.id)
                if not existing:
                    supabase.table("profiles").insert({
                        "id": session.user.id,
                        "email": session.user.email,
                        "plan": "free",
                        "credits_used_this_month": 0,
                        "credits_reset_at": datetime.now(timezone.utc).isoformat()
                    }).execute()
            except Exception:
                pass
            st.query_params.clear()
            for k in ["show_auth","show_login","show_register","guest_cv_data"]:
                st.session_state.pop(k, None)
            st.rerun()
    except Exception:
        pass

def handle_google_callback():
    """Intercept Google OAuth code from URL — processed exactly once."""
    if st.session_state.get("_google_handled"):
        return
    params = dict(st.query_params)
    code = params.get("code", "")
    if not code:
        return
    try:
        st.session_state["_google_handled"] = True  # mark BEFORE call to prevent reuse
        session = supabase.auth.exchange_code_for_session({"auth_code": code})
        if session and session.user:
            st.session_state["user"]    = session.user
            st.session_state["session"] = session
            try:
                existing = get_profile(session.user.id)
                if not existing:
                    supabase.table("profiles").insert({
                        "id": session.user.id,
                        "email": session.user.email,
                        "plan": "free",
                        "credits_used_this_month": 0,
                        "credits_reset_at": datetime.now(timezone.utc).isoformat()
                    }).execute()
            except Exception:
                pass
            st.query_params.clear()
            for k in ["show_auth","show_login","show_register","guest_cv_data"]:
                st.session_state.pop(k, None)
            st.rerun()
    except Exception as e:
        st.session_state.pop("_google_handled", None)  # allow retry on error

def get_google_oauth_url() -> str:
    """Get Google OAuth URL — redirects to Streamlit app with ?code= param."""
    try:
        res = supabase.auth.sign_in_with_oauth({
            "provider": "google",
            "options": {
                "redirect_to": "https://analyze-this-v2.streamlit.app",
                "skip_browser_redirect": True,
                "scopes": "openid email profile",
                "query_params": {"prompt": "select_account", "access_type": "offline"}
            }
        })
        return res.url if hasattr(res, "url") and res.url else ""
    except Exception:
        return ""

def sign_out():
    try:
        supabase.auth.sign_out()
    except Exception:
        pass
    for key in ["user","session","cv_data","regen_docx","api_credits_error"]:
        st.session_state.pop(key, None)
    st.rerun()

def send_password_reset(email: str) -> tuple[bool, str]:
    try:
        supabase.auth.reset_password_email(
            email,
            options={"redirect_to": "https://analyze-this-v2.streamlit.app"}
        )
        return True, "✅ Te enviamos un email con el link para restablecer tu contraseña."
    except Exception as e:
        return False, f"No se pudo enviar el email: {e}"

def get_profile(user_id: str) -> dict:
    try:
        res = supabase.table("profiles").select("*").eq("id", user_id).single().execute()
        return res.data or {}
    except Exception:
        return {}

def get_credits_remaining(profile: dict) -> int:
    plan = profile.get("plan", "free")
    if plan == "admin":
        return 999999
    monthly_limit = PLAN_CREDITS.get(plan, 5)
    used = profile.get("credits_used_this_month", 0)
    # Auto-reset if new month
    reset_at_str = profile.get("credits_reset_at", "")
    try:
        reset_at = datetime.fromisoformat(reset_at_str.replace("Z", "+00:00"))
        now = datetime.now(timezone.utc)
        if now.month != reset_at.month or now.year != reset_at.year:
            # Reset credits
            supabase.table("profiles").update({
                "credits_used_this_month": 0,
                "credits_reset_at": now.isoformat()
            }).eq("id", profile["id"]).execute()
            used = 0
    except Exception:
        pass
    return max(0, monthly_limit - used)

def consume_credit(user_id: str, current_used: int):
    supabase.table("profiles").update({
        "credits_used_this_month": current_used + 1
    }).eq("id", user_id).execute()

def save_history(user_id: str, job_title: str, score: int, ats_ok: bool) -> int:
    """Returns the new history row id so we can update outcome later."""
    try:
        res = supabase.table("history").insert({
            "user_id": user_id,
            "job_title": job_title[:120],
            "score_match": score,
            "ats_compatible": ats_ok,
            "outcome": None,  # null until user reports back
            "created_at": datetime.now(timezone.utc).isoformat()
        }).execute()
        if res.data:
            return res.data[0].get("id")
        return None
    except Exception:
        return None

def update_outcome(history_id: int, outcome: str):
    """User reports back: got_interview, got_job, no_response, rejected."""
    try:
        supabase.table("history").update({"outcome": outcome}).eq("id", history_id).execute()
        return True
    except Exception:
        return False

def save_cv_copy(user_id: str, history_id: int, cv_original: str, cv_data: dict):
    """Stores original CV text and generated JSON — opt-in only."""
    try:
        supabase.table("cv_storage").insert({
            "user_id": user_id,
            "history_id": history_id,
            "cv_original_snippet": cv_original[:5000],   # first 5k chars
            "cv_generated": json.dumps(cv_data, ensure_ascii=False)[:20000],
            "created_at": datetime.now(timezone.utc).isoformat()
        }).execute()
        return True
    except Exception:
        return False


def get_history(user_id: str) -> list:
    try:
        res = (supabase.table("history")
               .select("*")
               .eq("user_id", user_id)
               .order("created_at", desc=True)
               .limit(10)
               .execute())
        return res.data or []
    except Exception:
        return []

# ─── Feedback system ──────────────────────────────────────────────────────────
def save_feedback(user_id: str, email: str, rating: int, comment: str, job_title: str):
    try:
        supabase.table("feedback").insert({
            "user_id": user_id,
            "email": email,
            "rating": rating,
            "comment": comment[:500],
            "job_title": job_title[:120],
            "approved": False,  # admin must approve before showing publicly
            "created_at": datetime.now(timezone.utc).isoformat()
        }).execute()
        return True
    except Exception:
        return False

def get_public_reviews() -> list:
    try:
        res = (supabase.table("feedback")
               .select("rating,comment,job_title,created_at")
               .eq("approved", True)
               .order("created_at", desc=True)
               .limit(6)
               .execute())
        return res.data or []
    except Exception:
        return []

def get_all_feedback() -> list:
    """Admin only — returns all feedback including unapproved."""
    try:
        res = (supabase.table("feedback")
               .select("*")
               .order("created_at", desc=True)
               .limit(50)
               .execute())
        return res.data or []
    except Exception:
        return []

def approve_feedback(feedback_id: int, approve: bool):
    try:
        supabase.table("feedback").update({"approved": approve}).eq("id", feedback_id).execute()
        return True
    except Exception:
        return False

# ─── Usage counter ────────────────────────────────────────────────────────────
def get_global_stats() -> dict:
    """Returns total CVs generated (logged + guest) and registered users."""
    try:
        cvs   = supabase.table("history").select("id", count="exact").execute()
        guest = supabase.table("guest_analyses").select("id", count="exact").execute()
        users = supabase.table("profiles").select("id", count="exact").execute()
        total_cvs = (cvs.count or 0) + (guest.count or 0)
        return {
            "cvs":   total_cvs + COUNTER_BASE_CVS,
            "users": (users.count or 0) + COUNTER_BASE_USERS,
        }
    except Exception:
        return {"cvs": COUNTER_BASE_CVS, "users": COUNTER_BASE_USERS}

# ─── Activation codes ─────────────────────────────────────────────────────────
def validate_and_use_code(user_id: str, code: str) -> tuple[bool, str]:
    try:
        code = code.strip().upper()
        res = supabase.table("activation_codes").select("*").eq("code", code).single().execute()
        row = res.data
        if not row:
            return False, "Código inválido."
        if not row.get("active", True):
            return False, "Este código ya no está activo."
        max_uses = row.get("max_uses")
        used = row.get("uses_count", 0)
        if max_uses and used >= max_uses:
            return False, "Este código ya alcanzó su límite de usos."
        expires = row.get("expires_at")
        if expires:
            exp_dt = datetime.fromisoformat(expires.replace("Z", "+00:00"))
            if datetime.now(timezone.utc) > exp_dt:
                return False, "Este código ha expirado."
        plan = row.get("grants_plan", "pro_code")
        supabase.table("profiles").update({
            "plan": plan,
            "activation_code": code,
            "credits_used_this_month": 0
        }).eq("id", user_id).execute()
        supabase.table("activation_codes").update({
            "uses_count": used + 1
        }).eq("code", code).execute()
        label = PLAN_CREDITS.get(plan, 10)
        return True, f"✅ Código activado. Ahora tienes {label} análisis por mes."
    except Exception:
        return False, "Código inválido."

def get_all_codes() -> list:
    try:
        res = supabase.table("activation_codes").select("*").order("created_at", desc=True).execute()
        return res.data or []
    except Exception:
        return []

def create_code(code: str, description: str, max_uses: int, grants_plan: str, expires_at):
    try:
        supabase.table("activation_codes").insert({
            "code": code.strip().upper(),
            "description": description,
            "max_uses": max_uses if max_uses > 0 else None,
            "grants_plan": grants_plan,
            "uses_count": 0,
            "active": True,
            "expires_at": expires_at,
            "created_at": datetime.now(timezone.utc).isoformat()
        }).execute()
        return True
    except Exception:
        return False

# ─── Auth wall ────────────────────────────────────────────────────────────────
def show_auth_page():
    # Back to guest mode
    if st.session_state.get("guest_cv_data"):
        if st.button("← Volver al análisis", key="btn_back_to_guest"):
            for k in ["show_auth","show_login","show_register"]:
                st.session_state.pop(k, None)
            st.rerun()

    st.markdown(f"""
<div style="text-align:center;padding:1.8rem 0 1rem 0">
  <div style="font-size:2rem;margin-bottom:0.5rem">🎯</div>
  <h1 style="font-size:1.7rem;font-weight:800;margin:0;letter-spacing:-0.03em;color:var(--text)">
    Analyze-This <span class="app-beta">Beta</span>
  </h1>
  <p style="color:var(--text3);font-size:0.9rem;margin:0.5rem 0 0 0">
    Sube tu CV. Pega la oferta. Descarga listo para enviar.
  </p>
</div>
""", unsafe_allow_html=True)

    # Value prop — why register
    st.markdown("""
<div style="display:grid;grid-template-columns:1fr 1fr 1fr;gap:0.6rem;margin-bottom:1.2rem">
  <div class="feat-pill">
    <div class="feat-pill-icon">📄</div>
    <div class="feat-pill-title">5 análisis gratis</div>
    <div class="feat-pill-sub">al mes, sin tarjeta</div>
  </div>
  <div class="feat-pill">
    <div class="feat-pill-icon">🎨</div>
    <div class="feat-pill-title">4 templates</div>
    <div class="feat-pill-sub">listos para enviar</div>
  </div>
  <div class="feat-pill">
    <div class="feat-pill-icon">🎤</div>
    <div class="feat-pill-title">Coaching incluido</div>
    <div class="feat-pill-sub">carta + entrevista</div>
  </div>
</div>
""", unsafe_allow_html=True)

    st.markdown("---")
    tab_login, tab_magic, tab_signup = st.tabs(["🔑 Con contraseña", "✉️ Magic Link", "📝 Crear cuenta"])

    with tab_login:
        with st.form("form_login", clear_on_submit=False):
            email = st.text_input("Email", key="login_email",
                value=st.session_state.get("_remembered_email", ""))
            password = st.text_input("Contraseña", type="password", key="login_pw")
            submitted = st.form_submit_button("Entrar", use_container_width=True)
        if submitted:
            if not email or not password:
                st.error("Completa email y contraseña.")
            else:
                with st.spinner("Verificando..."):
                    ok, msg = sign_in(email, password)
                if ok:
                    st.session_state["_remembered_email"] = email
                    for k in ["show_auth","show_login","show_register","guest_cv_data"]:
                        st.session_state.pop(k, None)
                    st.rerun()
                else:
                    st.error(msg)

        # Forgot password
        with st.expander("¿Olvidaste tu contraseña?"):
            reset_email = st.text_input("Email de tu cuenta", key="reset_email")
            if st.button("Enviar link de recuperación", key="btn_reset"):
                if not reset_email:
                    st.error("Ingresa tu email.")
                else:
                    with st.spinner("Enviando..."):
                        ok, msg = send_password_reset(reset_email)
                    if ok:
                        st.success(msg)
                    else:
                        st.error(msg)

    with tab_magic:
        st.caption("Te enviamos un link a tu email — haz clic y entras directo, sin contraseña.")
        magic_email = st.text_input("Tu email", key="magic_email")
        if st.button("📩 Enviar link mágico", use_container_width=True, key="btn_magic"):
            if not magic_email:
                st.error("Ingresa tu email.")
            else:
                with st.spinner("Enviando..."):
                    ok, msg = send_magic_link(magic_email)
                if ok:
                    st.success(msg)
                    st.info("💡 El link expira en 1 hora. Si no llega, revisa spam.")
                else:
                    st.error(msg)

    with tab_signup:
        st.markdown("""
<div style="display:grid;grid-template-columns:1fr 1fr;gap:0.6rem;margin-bottom:0.9rem">
  <div class="at-card" style="text-align:center">
    <div style="font-size:0.72rem;font-weight:700;text-transform:uppercase;letter-spacing:0.08em;color:var(--text3)">Sin código</div>
    <div style="font-size:1.8rem;font-weight:800;color:var(--text);margin:0.2rem 0">5</div>
    <div style="font-size:0.75rem;color:var(--text3)">análisis / mes</div>
  </div>
  <div class="at-card" style="text-align:center;border-color:rgba(245,158,11,0.25)">
    <div style="font-size:0.72rem;font-weight:700;text-transform:uppercase;letter-spacing:0.08em;color:var(--amber)">Con código</div>
    <div style="font-size:1.8rem;font-weight:800;color:var(--text);margin:0.2rem 0">10</div>
    <div style="font-size:0.75rem;color:var(--text3)">análisis / mes</div>
  </div>
</div>
""", unsafe_allow_html=True)
        email2 = st.text_input("Email", key="signup_email")
        display_name = st.text_input("Nombre o apodo (opcional)",
            placeholder="Como quieres que te llamemos — ej: Rocío, Juan P.",
            key="signup_display_name")
        password2 = st.text_input("Contraseña (mín. 8 caracteres)", type="password", key="signup_pw")
        password3 = st.text_input("Confirmar contraseña", type="password", key="signup_pw2")
        activation_code = st.text_input("🎟️ Código de activación (opcional)",
            placeholder="Ingresa tu código si tienes uno — te da más análisis",
            key="signup_code")
        if st.button("Crear cuenta", use_container_width=True, key="btn_signup"):
            if not email2 or not password2:
                st.error("Completa todos los campos.")
            elif password2 != password3:
                st.error("Las contraseñas no coinciden.")
            elif len(password2) < 8:
                st.error("La contraseña debe tener al menos 8 caracteres.")
            else:
                with st.spinner("Creando cuenta..."):
                    ok, msg = sign_up(email2, password2, display_name.strip())
                if ok:
                    stage_val = None
                    found_val = None
                    if activation_code.strip():
                        import time; time.sleep(1.5)
                        try:
                            new_user = supabase.auth.sign_in_with_password({"email": email2, "password": password2})
                            if new_user and new_user.user:
                                # Save onboarding
                                try:
                                    supabase.table("profiles").update({
                                        "career_stage": stage_val, "how_found": found_val
                                    }).eq("id", new_user.user.id).execute()
                                except Exception: pass
                                code_ok, code_msg = validate_and_use_code(new_user.user.id, activation_code)
                                supabase.auth.sign_out()
                                if code_ok:
                                    st.success(f"¡Cuenta creada! {code_msg}")
                                else:
                                    st.success("¡Cuenta creada! (5 análisis/mes)")
                                    st.warning(f"Código: {code_msg}")
                            else:
                                st.success(msg)
                        except Exception:
                            st.success(msg)
                    else:
                        # No code — save onboarding briefly
                        try:
                            import time as _t; _t.sleep(1)
                            temp_u = supabase.auth.sign_in_with_password({"email": email2, "password": password2})
                            if temp_u and temp_u.user:
                                supabase.table("profiles").update({
                                    "career_stage": stage_val, "how_found": found_val
                                }).eq("id", temp_u.user.id).execute()
                                supabase.auth.sign_out()
                        except Exception: pass
                        st.success(msg)
                    st.info("👆 Inicia sesión en la pestaña **🔑 Iniciar sesión**.")
                else:
                    st.error(msg)

    st.markdown("---")
    st.markdown('<div class="section-label">Planes</div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("""<div class="at-card">
          <div style="font-size:0.8rem;font-weight:700;color:var(--text);margin-bottom:0.4rem">🆓 Free</div>
          <div style="font-size:0.75rem;color:var(--text2);line-height:1.7">5 análisis/mes<br>CV descargable<br>Análisis ATS + coaching</div>
        </div>""", unsafe_allow_html=True)
    with c2:
        st.markdown("""<div class="at-card" style="border-color:rgba(245,158,11,0.2)">
          <div style="font-size:0.8rem;font-weight:700;color:var(--amber);margin-bottom:0.4rem">⭐ Pro</div>
          <div style="font-size:0.75rem;color:var(--text2);line-height:1.7">50 análisis/mes<br>Todo lo de Free<br>Historial completo</div>
        </div>""", unsafe_allow_html=True)
    with c3:
        st.markdown("""<div class="at-card">
          <div style="font-size:0.8rem;font-weight:700;color:#c084fc;margin-bottom:0.4rem">🏢 Admin</div>
          <div style="font-size:0.75rem;color:var(--text2);line-height:1.7">Uso ilimitado<br>Panel de gestión<br>Vista de usuarios</div>
        </div>""", unsafe_allow_html=True)

    # Social proof counter
    stats = get_global_stats()
    if stats["cvs"] > 0:
        st.markdown("---")
        c1, c2 = st.columns(2)
        with c1:
            st.metric("📄 CVs optimizados", f"{stats['cvs']:,}")
        with c2:
            st.metric("👥 Usuarios registrados", f"{stats['users']:,}")

    st.markdown("---")
    st.markdown("""
<div style="text-align:center;padding:0.6rem 0">
  <a href="https://ko-fi.com/analyzethis" target="_blank"
     style="display:inline-block;background:#FFDD00;color:#000;font-weight:700;
     padding:0.45rem 1.2rem;border-radius:8px;text-decoration:none;font-size:0.85rem;">
    ☕ ¿Te fue útil? Apoya en Ko-fi
  </a>
  <p style="font-size:0.72rem;color:var(--text3);margin-top:0.4rem">
    Ayuda a mantener el servicio gratuito para todos
  </p>
</div>""", unsafe_allow_html=True)

    if not SUPABASE_URL or not SUPABASE_KEY:
        st.warning("⚠️ Supabase no configurado. Agrega SUPABASE_URL y SUPABASE_KEY en Secrets.")

# ─── Google OAuth via streamlit-oauth ────────────────────────────────────────
GOOGLE_AUTH_URL  = "https://accounts.google.com/o/oauth2/v2/auth"
GOOGLE_TOKEN_URL = "https://oauth2.googleapis.com/token"
GOOGLE_SCOPE     = "openid email profile"

def handle_google_token(token: dict):
    """Takes Google token, signs in to Supabase with the id_token."""
    try:
        id_token = token.get("id_token", "")
        if not id_token:
            st.error("No se pudo obtener el token de Google.")
            return
        # Sign in to Supabase with Google id_token
        res = supabase.auth.sign_in_with_id_token({
            "provider": "google",
            "token": id_token,
        })
        if res and res.user:
            st.session_state["user"] = res.user
            st.session_state["session"] = res.session
            # Create profile if first time
            try:
                existing = get_profile(res.user.id)
                if not existing:
                    supabase.table("profiles").insert({
                        "id": res.user.id,
                        "email": res.user.email,
                        "plan": "free",
                        "credits_used_this_month": 0,
                        "credits_reset_at": datetime.now(timezone.utc).isoformat()
                    }).execute()
            except Exception:
                pass
            st.rerun()
        else:
            st.error("Error al iniciar sesión con Google.")
    except Exception as e:
        st.error(f"Error de autenticación: {e}")

# ─── Extraction helpers ────────────────────────────────────────────────────────
def is_valid_url(url: str) -> bool:
    return bool(re.match(r'^https?://', url.strip()))

def scrape_job_url(url: str) -> str:
    headers = {"User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    )}
    try:
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
    except requests.exceptions.Timeout:
        raise ValueError("El sitio tardó demasiado. Pega el texto manualmente.")
    except requests.exceptions.HTTPError as e:
        raise ValueError(f"No se pudo acceder ({e.response.status_code}).")
    except Exception:
        raise ValueError("No se pudo acceder al link.")

    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script","style","nav","footer","header","aside","form","noscript","iframe"]):
        tag.decompose()
    text = ""
    for sel in [
        {"class": lambda c: c and any(k in " ".join(c).lower() for k in
            ["job-description","description","vacancy","posting","details","content"])},
        {"id": lambda i: i and any(k in i.lower() for k in ["job-description","description","details"])},
    ]:
        block = soup.find(attrs=sel)
        if block:
            text = block.get_text(separator="\n", strip=True)
            if len(text) > 200:
                break
    if not text or len(text) < 200:
        text = soup.get_text(separator="\n", strip=True)
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    return "\n".join(lines)[:8000]

def extract_pdf(file) -> str:
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    return text.strip()

def extract_docx(file) -> str:
    doc = DocxDocument(file)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

# ─── Claude optimization ──────────────────────────────────────────────────────
def optimize_cv(cv_text: str, job_text: str, max_pages: int, font_size, career_change: bool = False, cv_only: bool = False, output_lang: str = "es") -> dict:
    api_key = st.session_state.get("user_api_key") or ANTHROPIC_KEY
    limit = MAX_CV_CHARS_CAREER if career_change else MAX_CV_CHARS
    was_truncated = len(cv_text) > limit
    cv_text = cv_text[:limit]

    words_per_page = {9: 700, 10: 600, 10.5: 560, 11: 520, 12: 460}
    max_words = words_per_page.get(float(font_size), 580) * max_pages

    lang_map = {"es": "español", "en": "English", "pt": "português"}
    lang_instruction = f"\nIDIOMA DE SALIDA: Redacta TODO el CV optimizado (resumen, logros, habilidades, secciones) en {lang_map.get(output_lang, 'español')}. El coaching y análisis también en ese idioma. El CV original y la oferta pueden estar en idiomas distintos — eso no es problema, compara su contenido y genera el output en el idioma solicitado.\n"

    prompt = f"""Eres un experto coach de carrera y especialista en optimización de CVs para sistemas ATS.
{lang_instruction}
{"MODO CAMBIO DE CARRERA — incluye experiencia de TODOS los períodos y reenmarca habilidades transferibles hacia el nuevo rol." if career_change else ("MODO ANÁLISIS GENERAL — no hay oferta específica. Analiza el CV: claridad ATS, modernidad, redacción de logros, estructura. Da recomendaciones concretas de mejora." if cv_only else "Selecciona la experiencia MÁS RECIENTE y relevante para esta oferta específica.")}

REGLA ABSOLUTA — CERO INVENCIÓN:
- Cada skill, herramienta, cargo, certificación y logro DEBE existir en el CV original
- Si el CV no menciona SAP → NO pongas SAP
- Si el CV no dice "Ingeniero de Materiales" → NO pongas ese título
- Si no hay experiencia en abastecimiento → NO pongas "especialista en abastecimiento"
- El título profesional debe ser el cargo real del CV más reciente, adaptado a la oferta — nunca inventado
- Ante la duda entre incluir algo o no → NO lo incluyas
Inventar datos es el error más grave posible. Es preferible un CV más corto que uno con datos falsos.

INSTRUCCIONES:
1. Solo reorganiza y reescribe lo que ya existe — NUNCA agregues información nueva
2. Integra las palabras clave de la oferta SOLO si tienen respaldo en el CV original
3. Reescribe logros con verbos de acción — pero con hechos reales del CV
4. Respeta el límite de {{max_pages}} página(s) ≈ {{max_words}} palabras
5. Detecta qué ATS probablemente usa la empresa según su nombre/industria:
   - Empresas grandes/corporativas → Workday, SAP SuccessFactors
   - Startups/tech latam → Greenhouse, Lever
   - Empresas chilenas → Buk, Rankmi, Talently
   - Retail/gobierno → sistemas propios básicos
   Adapta el formato y densidad de keywords según ese ATS inferido
6. Evalúa compatibilidad ATS: sin tablas ni columnas que confundan parsers

CV ORIGINAL:
{cv_text}

OFERTA DE TRABAJO:
{job_text}

Responde ÚNICAMENTE con JSON válido, sin backticks:
{{
  "nombre": "nombre completo",
  "titulo_profesional": "título adaptado al puesto",
  "email": "email o vacío",
  "telefono": "teléfono o vacío",
  "linkedin": "URL o vacío",
  "ubicacion": "ciudad, país",
  "resumen_profesional": "3-4 oraciones con keywords ATS.",
  "experiencia": [{{
    "empresa": "nombre",
    "cargo": "cargo",
    "periodo": "mes/año - mes/año",
    "logros": ["Logro cuantificado con keyword ATS"]
  }}],
  "educacion": [{{"institucion":"","titulo":"","periodo":"","detalle":""}}],
  "habilidades_tecnicas": ["skill relevante"],
  "habilidades_blandas": ["máx 4"],
  "idiomas": ["Español - Nativo"],
  "certificaciones": ["solo si existen"],
  "ats_compatible": true,
  "ats_detectado": "Nombre del ATS inferido para esta empresa (ej: Workday, Buk, Greenhouse)",
  "ats_razon": "Por qué se adaptó el CV para ese ATS específico",
  "score_match": 82,
  "score_desglose": {{"keywords":88,"experiencia":80,"educacion":75,"habilidades":85}},
  "score_explicacion": "2-3 oraciones sobre el score.",
  "keywords_integradas": ["kw1"],
  "keywords_faltantes": ["kw ausente"],
  "coaching": [
    {{"categoria": "Tu fortaleza clave 💪", "tip": "Qué tiene el candidato valioso y cómo destacarlo."}},
    {{"categoria": "Brecha crítica ⚠️", "tip": "Skill que falta y cómo cerrarla con curso/cert específica."}},
    {{"categoria": "Quick win de hoy ⚡", "tip": "Acción concreta en menos de 1 hora para mejorar candidatura."}},
    {{"categoria": "LinkedIn / Marca personal 🔗", "tip": "Qué cambiar en LinkedIn para este rol."}},
    {{"categoria": "Antes de la entrevista 📋", "tip": "Qué investigar y qué narrativa preparar."}}
  ]
}}"""

    client = anthropic.Anthropic(api_key=api_key)

    def call_model(model_id: str) -> dict:
        for attempt in range(2):
            msg = client.messages.create(
                model=model_id,
                max_tokens=5000,
                temperature=0,
                messages=[{"role": "user", "content": prompt}]
            )
            raw = msg.content[0].text.strip()
            if "```json" in raw:
                raw = raw.split("```json")[1].split("```")[0].strip()
            elif "```" in raw:
                raw = raw.split("```")[1].split("```")[0].strip()
            try:
                return json.loads(raw)
            except json.JSONDecodeError:
                if attempt == 0:
                    time.sleep(1)
                    continue
                raise

    # TODO future: add style variation pass to reduce AI-detectable patterns
    # (vary sentence length, avoid Claude's typical connectors, add burstiness)
    # Not urgent until market starts flagging AI-written CVs in Latam.

    # Paso 1: Haiku hace el análisis inicial (costo ~$0.01-0.02)
    result = call_model("claude-haiku-4-5-20251001")
    result["_was_truncated"] = was_truncated
    result["_model_used"] = "haiku"

    # Paso 2: Si score < 60, Opus reanaliza con más profundidad (~$0.15-0.20)
    # Solo ocurre cuando el candidato realmente necesita optimización más inteligente
    if result.get("score_match", 100) < 60:
        result = call_model("claude-opus-4-5-20251101")
        result["_was_truncated"] = was_truncated
        result["_model_used"] = "opus"

    return result

# ─── DOCX helpers ─────────────────────────────────────────────────────────────
def add_section_header(doc, title, color_rgb, fn, body_size, border_color, prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(body_size)
    p.paragraph_format.space_after = Pt(body_size * 0.3)
    run = p.add_run(f"{prefix}{title}")
    run.bold = True
    run.font.name = fn
    run.font.size = Pt(body_size + 1)
    run.font.color.rgb = RGBColor(*color_rgb)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.append(pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), border_color)
    pBdr.append(bottom)

def R(run, fn, fs, bold=False, italic=False, color=None):
    run.font.name = fn
    run.font.size = Pt(float(fs))
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def build_classic(cv, fn, fs):
    fs = float(fs)
    doc = DocxDocument()
    s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(1.0)
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)
    DARK=(0x1A,0x1A,0x2E); BLUE=(0x2E,0x75,0xB6); GRAY=(0x66,0x66,0x66)
    def hdr(t): add_section_header(doc, t, BLUE, fn, fs, "2E75B6")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    R(p.add_run(cv.get("nombre","")), fn, fs+9, bold=True, color=DARK)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    R(p.add_run(cv.get("titulo_profesional","")), fn, fs+2, bold=True, color=BLUE)
    parts=[x for x in [cv.get("email"),cv.get("telefono"),cv.get("ubicacion"),cv.get("linkedin")] if x]
    if parts:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        R(p.add_run("  |  ".join(parts)), fn, fs-1, color=GRAY)
    if cv.get("resumen_profesional"):
        hdr("RESUMEN PROFESIONAL")
        R(doc.add_paragraph().add_run(cv["resumen_profesional"]), fn, fs)
    if cv.get("experiencia"):
        hdr("EXPERIENCIA PROFESIONAL")
        for exp in cv["experiencia"]:
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(fs*0.6)
            R(p.add_run(exp.get("cargo","")), fn, fs, bold=True)
            p2=doc.add_paragraph()
            R(p2.add_run(f"{exp.get('empresa','')}   |   {exp.get('periodo','')}"), fn, fs-1, italic=True, color=GRAY)
            for logro in exp.get("logros",[]):
                pb=doc.add_paragraph(style="List Bullet")
                pb.paragraph_format.left_indent=Inches(0.2); pb.paragraph_format.space_after=Pt(2)
                R(pb.add_run(logro), fn, fs)
    if cv.get("educacion"):
        hdr("EDUCACIÓN")
        for edu in cv["educacion"]:
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(fs*0.5)
            R(p.add_run(edu.get("titulo","")), fn, fs, bold=True)
            p2=doc.add_paragraph()
            R(p2.add_run(f"{edu.get('institucion','')}   |   {edu.get('periodo','')}"), fn, fs-1, italic=True, color=GRAY)
            if edu.get("detalle"): R(doc.add_paragraph().add_run(edu["detalle"]), fn, fs-1)
    for label, key, header in [
        (None,"habilidades_tecnicas","HABILIDADES TÉCNICAS"),
        (None,"habilidades_blandas","COMPETENCIAS"),
    ]:
        if cv.get(key):
            hdr(header)
            R(doc.add_paragraph().add_run("  •  ".join(cv[key])), fn, fs)
    if cv.get("idiomas"):
        hdr("IDIOMAS"); R(doc.add_paragraph().add_run("  |  ".join(cv["idiomas"])), fn, fs)
    certs=[c for c in cv.get("certificaciones",[]) if c]
    if certs:
        hdr("CERTIFICACIONES")
        for cert in certs: R(doc.add_paragraph().add_run(f"• {cert}"), fn, fs)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf

def build_modern(cv, fn, fs):
    fs = float(fs)
    doc = DocxDocument()
    s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(0.75)
    s.top_margin = Inches(0.6)
    s.bottom_margin = Inches(0.8)
    NAVY=(0x1B,0x4F,0x72); TEAL=(0x17,0x8A,0xCA); GRAY=(0x77,0x77,0x77)
    def hdr(t): add_section_header(doc, t, NAVY, fn, fs, "17A8CA", prefix="◆  ")
    R(doc.add_paragraph().add_run(cv.get("nombre","").upper()), fn, fs+11, bold=True, color=NAVY)
    R(doc.add_paragraph().add_run(cv.get("titulo_profesional","")), fn, fs+2, bold=True, color=TEAL)
    parts=[]
    for icon,key in [("✉","email"),("✆","telefono"),("⌖","ubicacion"),("in","linkedin")]:
        if cv.get(key): parts.append(f"{icon} {cv[key]}")
    if parts: R(doc.add_paragraph().add_run("   |   ".join(parts)), fn, fs-1, color=GRAY)
    p_div=doc.add_paragraph(); p_div.paragraph_format.space_after=Pt(8)
    pPr=p_div._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr'); pPr.append(pBdr)
    btm=OxmlElement('w:bottom')
    for a,v in [('w:val','single'),('w:sz','16'),('w:space','1'),('w:color','1B4F72')]: btm.set(qn(a),v)
    pBdr.append(btm)
    if cv.get("resumen_profesional"):
        hdr("PERFIL PROFESIONAL")
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.15)
        R(p.add_run(cv["resumen_profesional"]), fn, fs)
    if cv.get("experiencia"):
        hdr("EXPERIENCIA")
        for exp in cv["experiencia"]:
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(fs*0.7); p.paragraph_format.left_indent=Inches(0.15)
            R(p.add_run(exp.get("cargo","")), fn, fs, bold=True, color=NAVY)
            R(p.add_run("  —  "), fn, fs)
            R(p.add_run(exp.get("empresa","")), fn, fs)
            p2=doc.add_paragraph(); p2.paragraph_format.left_indent=Inches(0.15)
            R(p2.add_run(exp.get("periodo","")), fn, fs-1, italic=True, color=TEAL)
            for logro in exp.get("logros",[]):
                pb=doc.add_paragraph(); pb.paragraph_format.left_indent=Inches(0.35); pb.paragraph_format.space_after=Pt(2)
                R(pb.add_run(f"▸  {logro}"), fn, fs)
    if cv.get("educacion"):
        hdr("EDUCACIÓN")
        for edu in cv["educacion"]:
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.15); p.paragraph_format.space_before=Pt(fs*0.5)
            R(p.add_run(edu.get("titulo","")), fn, fs, bold=True, color=NAVY)
            R(p.add_run(f"   |   {edu.get('institucion','')}   |   {edu.get('periodo','')}"), fn, fs-1)
            if edu.get("detalle"):
                p2=doc.add_paragraph(); p2.paragraph_format.left_indent=Inches(0.15)
                R(p2.add_run(edu["detalle"]), fn, fs-1)
    if cv.get("habilidades_tecnicas") or cv.get("habilidades_blandas"):
        hdr("HABILIDADES")
        for label,key in [("Técnicas: ","habilidades_tecnicas"),("Competencias: ","habilidades_blandas")]:
            if cv.get(key):
                p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.15)
                R(p.add_run(label), fn, fs, bold=True); R(p.add_run("  •  ".join(cv[key])), fn, fs)
    if cv.get("idiomas"):
        hdr("IDIOMAS")
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.15)
        R(p.add_run("  |  ".join(cv["idiomas"])), fn, fs)
    certs=[c for c in cv.get("certificaciones",[]) if c]
    if certs:
        hdr("CERTIFICACIONES")
        for cert in certs:
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.15)
            R(p.add_run(f"▸  {cert}"), fn, fs)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf

# ─── Template definitions ────────────────────────────────────────────────────
TEMPLATES = {
    "Clásico": {
        "icon": "📋",
        "color": "#2E75B6",
        "ideal": "Finanzas · Legal · Gobierno · Roles senior",
        "desc": "Nombre centrado, secciones con línea azul, bullets ordenados. El formato más reconocido y esperado por reclutadores tradicionales.",
        "tags": ["#EBF3FB", "#2E75B6"],
    },
    "Moderno": {
        "icon": "✨",
        "color": "#178ACA",
        "ideal": "Tech · Startups · Marketing · Diseño",
        "desc": "Header con nombre en mayúsculas, secciones con rombos ◆ y flechas ▸. Diseño limpio que destaca sin sacrificar legibilidad ATS.",
        "tags": ["#E8F7FD", "#178ACA"],
    },
    "Ejecutivo": {
        "icon": "🏛️",
        "color": "#1B2A4A",
        "ideal": "Dirección general · C-level · Consultoría · Banca",
        "desc": "Nombre en mayúsculas con título en dorado, separador ancho, tipografía Georgia. Transmite autoridad y trayectoria desde la primera línea.",
        "tags": ["#EDEEF2", "#1B2A4A"],
    },
    "Minimalista": {
        "icon": "⬜",
        "color": "#444444",
        "ideal": "Cualquier sector · Máxima compatibilidad ATS",
        "desc": "Sin colores, sin elementos gráficos. El formato más seguro para sistemas ATS estrictos. Ideal si no sabes qué parser usa la empresa.",
        "tags": ["#F5F5F5", "#444444"],
    },
}

from docx.oxml.ns import qn as _qn
from docx.oxml import OxmlElement as _OxmlElement

def _section_border(p, border_color):
    pPr = p._p.get_or_add_pPr()
    pBdr = _OxmlElement('w:pBdr')
    pPr.append(pBdr)
    bottom = _OxmlElement('w:bottom')
    bottom.set(_qn('w:val'), 'single')
    bottom.set(_qn('w:sz'), '6')
    bottom.set(_qn('w:space'), '1')
    bottom.set(_qn('w:color'), border_color)
    pBdr.append(bottom)

def _hdr(doc, title, color_rgb, fn, fs, border_color, prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(fs)
    p.paragraph_format.space_after = Pt(fs * 0.3)
    run = p.add_run(f"{prefix}{title}")
    run.bold = True; run.font.name = fn
    run.font.size = Pt(fs + 1)
    run.font.color.rgb = RGBColor(*color_rgb)
    _section_border(p, border_color)

def _R(run, fn, fs, bold=False, italic=False, color=None):
    run.font.name = fn; run.font.size = Pt(float(fs))
    run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = RGBColor(*color)
    return run

def _exp_block(doc, cv, fn, fs, cargo_color, periodo_color, bullet_prefix, indent_cargo, indent_bullet):
    for exp in cv.get("experiencia", []):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(fs * 0.6)
        p.paragraph_format.left_indent = Inches(indent_cargo)
        _R(p.add_run(exp.get("cargo", "")), fn, fs, bold=True, color=cargo_color)
        p2 = doc.add_paragraph(); p2.paragraph_format.left_indent = Inches(indent_cargo)
        _R(p2.add_run(f"{exp.get('empresa','')}   |   {exp.get('periodo','')}"), fn, fs-1, italic=True, color=periodo_color)
        for logro in exp.get("logros", []):
            pb = doc.add_paragraph(); pb.paragraph_format.left_indent = Inches(indent_bullet)
            pb.paragraph_format.space_after = Pt(2)
            _R(pb.add_run(f"{bullet_prefix}{logro}"), fn, fs)

def build_clasico(cv, fn, fs):
    fs = float(fs)
    doc = DocxDocument(); s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(1.0)
    s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)
    DARK=(0x1A,0x1A,0x2E); BLUE=(0x2E,0x75,0xB6); GRAY=(0x66,0x66,0x66)
    def hdr(t): _hdr(doc, t, BLUE, fn, fs, "2E75B6")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _R(p.add_run(cv.get("nombre","")), fn, fs+9, bold=True, color=DARK)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _R(p.add_run(cv.get("titulo_profesional","")), fn, fs+2, bold=True, color=BLUE)
    parts=[x for x in [cv.get("email"),cv.get("telefono"),cv.get("ubicacion"),cv.get("linkedin")] if x]
    if parts:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _R(p.add_run("  |  ".join(parts)), fn, fs-1, color=GRAY)
    if cv.get("resumen_profesional"):
        hdr("RESUMEN PROFESIONAL"); _R(doc.add_paragraph().add_run(cv["resumen_profesional"]), fn, fs)
    if cv.get("experiencia"):
        hdr("EXPERIENCIA PROFESIONAL")
        _exp_block(doc, cv, fn, fs, (0,0,0), GRAY, "• ", 0, 0.2)
    if cv.get("educacion"):
        hdr("EDUCACIÓN")
        for edu in cv["educacion"]:
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(fs*0.5)
            _R(p.add_run(edu.get("titulo","")), fn, fs, bold=True)
            _R(doc.add_paragraph().add_run(f"{edu.get('institucion','')}   |   {edu.get('periodo','')}"), fn, fs-1, italic=True, color=GRAY)
            if edu.get("detalle"): _R(doc.add_paragraph().add_run(edu["detalle"]), fn, fs-1)
    if cv.get("habilidades_tecnicas"):
        hdr("HABILIDADES TÉCNICAS"); _R(doc.add_paragraph().add_run("  •  ".join(cv["habilidades_tecnicas"])), fn, fs)
    if cv.get("habilidades_blandas"):
        hdr("COMPETENCIAS"); _R(doc.add_paragraph().add_run("  •  ".join(cv["habilidades_blandas"])), fn, fs)
    if cv.get("idiomas"):
        hdr("IDIOMAS"); _R(doc.add_paragraph().add_run("  |  ".join(cv["idiomas"])), fn, fs)
    certs=[c for c in cv.get("certificaciones",[]) if c]
    if certs:
        hdr("CERTIFICACIONES")
        for cert in certs: _R(doc.add_paragraph().add_run(f"• {cert}"), fn, fs)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf

def build_moderno(cv, fn, fs):
    fs = float(fs)
    doc = DocxDocument(); s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(0.75)
    s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.8)
    NAVY=(0x1B,0x4F,0x72); TEAL=(0x17,0x8A,0xCA); GRAY=(0x77,0x77,0x77)
    def hdr(t): _hdr(doc, t, NAVY, fn, fs, "17A8CA", prefix="◆  ")
    _R(doc.add_paragraph().add_run(cv.get("nombre","").upper()), fn, fs+11, bold=True, color=NAVY)
    _R(doc.add_paragraph().add_run(cv.get("titulo_profesional","")), fn, fs+2, bold=True, color=TEAL)
    parts=[]
    for icon,key in [("✉","email"),("✆","telefono"),("⌖","ubicacion"),("in","linkedin")]:
        if cv.get(key): parts.append(f"{icon} {cv[key]}")
    if parts: _R(doc.add_paragraph().add_run("   |   ".join(parts)), fn, fs-1, color=GRAY)
    p_div=doc.add_paragraph(); p_div.paragraph_format.space_after=Pt(8)
    pPr=p_div._p.get_or_add_pPr(); pBdr=_OxmlElement('w:pBdr'); pPr.append(pBdr)
    btm=_OxmlElement('w:bottom')
    for a,v in [('w:val','single'),('w:sz','16'),('w:space','1'),('w:color','1B4F72')]: btm.set(_qn(a),v)
    pBdr.append(btm)
    if cv.get("resumen_profesional"):
        hdr("PERFIL PROFESIONAL")
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.15)
        _R(p.add_run(cv["resumen_profesional"]), fn, fs)
    if cv.get("experiencia"):
        hdr("EXPERIENCIA")
        for exp in cv["experiencia"]:
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(fs*0.7); p.paragraph_format.left_indent=Inches(0.15)
            _R(p.add_run(exp.get("cargo","")), fn, fs, bold=True, color=NAVY)
            _R(p.add_run("  —  "), fn, fs); _R(p.add_run(exp.get("empresa","")), fn, fs)
            p2=doc.add_paragraph(); p2.paragraph_format.left_indent=Inches(0.15)
            _R(p2.add_run(exp.get("periodo","")), fn, fs-1, italic=True, color=TEAL)
            for logro in exp.get("logros",[]):
                pb=doc.add_paragraph(); pb.paragraph_format.left_indent=Inches(0.35); pb.paragraph_format.space_after=Pt(2)
                _R(pb.add_run(f"▸  {logro}"), fn, fs)
    if cv.get("educacion"):
        hdr("EDUCACIÓN")
        for edu in cv["educacion"]:
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.15); p.paragraph_format.space_before=Pt(fs*0.5)
            _R(p.add_run(edu.get("titulo","")), fn, fs, bold=True, color=NAVY)
            _R(p.add_run(f"   |   {edu.get('institucion','')}   |   {edu.get('periodo','')}"), fn, fs-1)
            if edu.get("detalle"):
                p2=doc.add_paragraph(); p2.paragraph_format.left_indent=Inches(0.15); _R(p2.add_run(edu["detalle"]), fn, fs-1)
    if cv.get("habilidades_tecnicas") or cv.get("habilidades_blandas"):
        hdr("HABILIDADES")
        for label,key in [("Técnicas: ","habilidades_tecnicas"),("Competencias: ","habilidades_blandas")]:
            if cv.get(key):
                p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.15)
                _R(p.add_run(label), fn, fs, bold=True); _R(p.add_run("  •  ".join(cv[key])), fn, fs)
    if cv.get("idiomas"):
        hdr("IDIOMAS"); p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.15)
        _R(p.add_run("  |  ".join(cv["idiomas"])), fn, fs)
    certs=[c for c in cv.get("certificaciones",[]) if c]
    if certs:
        hdr("CERTIFICACIONES")
        for cert in certs:
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.15); _R(p.add_run(f"▸  {cert}"), fn, fs)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf

def build_ejecutivo(cv, fn, fs):
    fs = float(fs)
    sfn = "Georgia" if fn in ["Calibri","Arial","Trebuchet MS"] else fn
    doc = DocxDocument(); s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(1.0)
    s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.8)
    NAVY=(0x1B,0x2A,0x4A); GOLD=(0x8B,0x6C,0x1E); GRAY=(0x55,0x55,0x55)
    def hdr(t): _hdr(doc, t, NAVY, sfn, fs, "1B2A4A")
    p=doc.add_paragraph(); _R(p.add_run(cv.get("nombre","").upper()), sfn, fs+10, bold=True, color=NAVY)
    p2=doc.add_paragraph(); _R(p2.add_run(cv.get("titulo_profesional","")), sfn, fs+1, italic=True, color=GOLD)
    parts=[x for x in [cv.get("email"),cv.get("telefono"),cv.get("ubicacion"),cv.get("linkedin")] if x]
    if parts: _R(doc.add_paragraph().add_run("  ·  ".join(parts)), fn, fs-1, color=GRAY)
    p_div=doc.add_paragraph(); p_div.paragraph_format.space_after=Pt(6)
    pPr=p_div._p.get_or_add_pPr(); pBdr=_OxmlElement('w:pBdr'); pPr.append(pBdr)
    btm=_OxmlElement('w:bottom')
    for a,v in [('w:val','single'),('w:sz','24'),('w:space','1'),('w:color','1B2A4A')]: btm.set(_qn(a),v)
    pBdr.append(btm)
    if cv.get("resumen_profesional"):
        hdr("PERFIL EJECUTIVO"); _R(doc.add_paragraph().add_run(cv["resumen_profesional"]), sfn, fs, italic=True)
    if cv.get("experiencia"):
        hdr("TRAYECTORIA PROFESIONAL")
        for exp in cv["experiencia"]:
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(fs*0.7)
            _R(p.add_run(f"■  {exp.get('cargo','')}"), sfn, fs, bold=True, color=NAVY)
            _R(p.add_run(f"  ·  {exp.get('empresa','')}"), sfn, fs)
            p2=doc.add_paragraph(); p2.paragraph_format.left_indent=Inches(0.25)
            _R(p2.add_run(exp.get("periodo","")), fn, fs-1, italic=True, color=GRAY)
            for logro in exp.get("logros",[]):
                pb=doc.add_paragraph(); pb.paragraph_format.left_indent=Inches(0.35); pb.paragraph_format.space_after=Pt(2)
                _R(pb.add_run(f"›  {logro}"), sfn, fs)
    if cv.get("educacion"):
        hdr("FORMACIÓN ACADÉMICA")
        for edu in cv["educacion"]:
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(fs*0.5)
            _R(p.add_run(edu.get("titulo","")), sfn, fs, bold=True, color=NAVY)
            _R(p.add_run(f"  —  {edu.get('institucion','')}  |  {edu.get('periodo','')}"), sfn, fs-1, color=GRAY)
    all_sk=cv.get("habilidades_tecnicas",[])+cv.get("habilidades_blandas",[])
    if all_sk:
        hdr("COMPETENCIAS"); _R(doc.add_paragraph().add_run("  ■  ".join(all_sk)), sfn, fs)
    if cv.get("idiomas"):
        hdr("IDIOMAS"); _R(doc.add_paragraph().add_run("  |  ".join(cv["idiomas"])), sfn, fs)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf

def build_minimalista(cv, fn, fs):
    fs = float(fs)
    doc = DocxDocument(); s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(1.1)
    s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
    DARK=(0x22,0x22,0x22); MID=(0x44,0x44,0x44); LIGHT=(0x88,0x88,0x88)
    def hdr(t):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(fs*1.2); p.paragraph_format.space_after=Pt(fs*0.2)
        run=p.add_run(t.upper()); run.bold=True; run.font.name=fn; run.font.size=Pt(fs); run.font.color.rgb=RGBColor(*DARK)
        _section_border(p, "AAAAAA")
    _R(doc.add_paragraph().add_run(cv.get("nombre","")), fn, fs+6, bold=True, color=DARK)
    _R(doc.add_paragraph().add_run(cv.get("titulo_profesional","")), fn, fs+0.5, color=MID)
    parts=[x for x in [cv.get("email"),cv.get("telefono"),cv.get("ubicacion"),cv.get("linkedin")] if x]
    if parts: _R(doc.add_paragraph().add_run("  |  ".join(parts)), fn, fs-1, color=LIGHT)
    if cv.get("resumen_profesional"):
        hdr("Resumen"); _R(doc.add_paragraph().add_run(cv["resumen_profesional"]), fn, fs, color=MID)
    if cv.get("experiencia"):
        hdr("Experiencia")
        for exp in cv["experiencia"]:
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(fs*0.5)
            _R(p.add_run(exp.get("cargo","")), fn, fs, bold=True, color=DARK)
            _R(p.add_run(f"  —  {exp.get('empresa','')}"), fn, fs, color=MID)
            _R(doc.add_paragraph().add_run(exp.get("periodo","")), fn, fs-1, italic=True, color=LIGHT)
            for logro in exp.get("logros",[]):
                pb=doc.add_paragraph(); pb.paragraph_format.left_indent=Inches(0.2); pb.paragraph_format.space_after=Pt(2)
                _R(pb.add_run(f"- {logro}"), fn, fs, color=MID)
    if cv.get("educacion"):
        hdr("Educación")
        for edu in cv["educacion"]:
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(fs*0.4)
            _R(p.add_run(edu.get("titulo","")), fn, fs, bold=True, color=DARK)
            _R(p.add_run(f"  —  {edu.get('institucion','')}  ({edu.get('periodo','')}"), fn, fs-1, color=MID)
    all_sk=cv.get("habilidades_tecnicas",[])+cv.get("habilidades_blandas",[])
    if all_sk:
        hdr("Habilidades"); _R(doc.add_paragraph().add_run("  /  ".join(all_sk)), fn, fs, color=MID)
    if cv.get("idiomas"):
        hdr("Idiomas"); _R(doc.add_paragraph().add_run("  |  ".join(cv["idiomas"])), fn, fs, color=MID)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf

BUILDERS = {
    "Clásico":    build_clasico,
    "Moderno":    build_moderno,
    "Ejecutivo":  build_ejecutivo,
    "Minimalista": build_minimalista,
}

# ─── CV PDF builder (reportlab) ───────────────────────────────────────────────
DISCLAIMER_TEXT = (
    "Este documento fue generado automáticamente por Analyze-This · CV Optimizer ATS "
    "usando inteligencia artificial (Claude, Anthropic). El contenido se basa "
    "exclusivamente en la información provista por el usuario — la herramienta reorganiza "
    "y optimiza, pero no verifica ni valida los datos ingresados. "
    "El usuario es el único responsable de la exactitud de su CV. "
    "analyze-this-v2.streamlit.app"
)

def build_cv_pdf(cv: dict, template: str = "Clásico") -> io.BytesIO:
    """Generate a clean PDF CV using reportlab — same content as DOCX templates."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors as _rc
    from reportlab.lib.units import cm
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     HRFlowable, KeepTogether)
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

    # Color schemes per template
    schemes = {
        "Clásico":    {"h": "#2E75B6", "t": "#1A1A2E", "s": "#666666", "acc": "#2E75B6"},
        "Moderno":    {"h": "#1B4F72", "t": "#1B4F72", "s": "#777777", "acc": "#178ACA"},
        "Ejecutivo":  {"h": "#1B2A4A", "t": "#1B2A4A", "s": "#555555", "acc": "#8B6C1E"},
        "Minimalista":{"h": "#222222", "t": "#222222", "s": "#888888", "acc": "#444444"},
    }
    sc = schemes.get(template, schemes["Clásico"])

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm, topMargin=1.8*cm, bottomMargin=2*cm)

    C_H   = _rc.HexColor(sc["h"])
    C_T   = _rc.HexColor(sc["t"])
    C_S   = _rc.HexColor(sc["s"])
    C_ACC = _rc.HexColor(sc["acc"])
    C_DIS = _rc.HexColor("#AAAAAA")

    fn = "Helvetica"

    def sty(name, **kw):
        return ParagraphStyle(name, fontName=fn, **kw)

    s_name   = sty("nm",  fontSize=18, fontName="Helvetica-Bold", textColor=C_T,
                   alignment=TA_CENTER, spaceAfter=3)
    s_title  = sty("tit", fontSize=11, fontName="Helvetica-Bold", textColor=C_H,
                   alignment=TA_CENTER, spaceAfter=3)
    s_contact= sty("con", fontSize=8.5, textColor=C_S,
                   alignment=TA_CENTER, spaceAfter=6)
    s_sec    = sty("sec", fontSize=9.5, fontName="Helvetica-Bold", textColor=C_H,
                   spaceBefore=10, spaceAfter=2)
    s_body   = sty("bod", fontSize=9, textColor=C_T, leading=14, spaceAfter=2)
    s_italic = sty("ita", fontSize=8.5, textColor=C_S, leading=13, spaceAfter=1)
    s_bullet = sty("bul", fontSize=9, textColor=C_T, leading=13,
                   leftIndent=12, bulletIndent=0, spaceAfter=1)
    s_disc   = sty("dis", fontSize=7, textColor=C_DIS, leading=10,
                   alignment=TA_CENTER, spaceBefore=8)

    def hr(color=C_H, thickness=0.8):
        return HRFlowable(width="100%", thickness=thickness,
                          color=color, spaceAfter=4)

    story = []

    # Header
    story.append(Paragraph(cv.get("nombre", ""), s_name))
    if cv.get("titulo_profesional"):
        story.append(Paragraph(cv["titulo_profesional"], s_title))
    parts = [x for x in [cv.get("email"), cv.get("telefono"),
                          cv.get("ubicacion"), cv.get("linkedin")] if x]
    if parts:
        story.append(Paragraph("  |  ".join(parts), s_contact))
    story.append(hr())

    def section(title):
        story.append(Paragraph(title.upper(), s_sec))
        story.append(hr(C_ACC, 0.5))

    if cv.get("resumen_profesional"):
        section("Resumen profesional")
        story.append(Paragraph(cv["resumen_profesional"], s_body))

    if cv.get("experiencia"):
        section("Experiencia profesional")
        for exp in cv["experiencia"]:
            block = []
            block.append(Paragraph(
                f"<b>{exp.get('cargo','')}</b>  —  {exp.get('empresa','')}",
                s_body))
            block.append(Paragraph(exp.get("periodo",""), s_italic))
            for logro in exp.get("logros", []):
                safe = logro.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
                block.append(Paragraph(f"• {safe}", s_bullet))
            story.append(KeepTogether(block))
            story.append(Spacer(1, 4))

    if cv.get("educacion"):
        section("Educación")
        for edu in cv["educacion"]:
            story.append(Paragraph(
                f"<b>{edu.get('titulo','')}</b>  —  {edu.get('institucion','')}  "
                f"({edu.get('periodo','')})", s_body))
            if edu.get("detalle"):
                story.append(Paragraph(edu["detalle"], s_italic))

    sk_tec = cv.get("habilidades_tecnicas", [])
    sk_bla = cv.get("habilidades_blandas", [])
    if sk_tec or sk_bla:
        section("Habilidades")
        if sk_tec:
            story.append(Paragraph(
                "<b>Técnicas:</b>  " + "  ·  ".join(sk_tec), s_body))
        if sk_bla:
            story.append(Paragraph(
                "<b>Competencias:</b>  " + "  ·  ".join(sk_bla), s_body))

    if cv.get("idiomas"):
        section("Idiomas")
        story.append(Paragraph("  |  ".join(cv["idiomas"]), s_body))

    certs = [c for c in cv.get("certificaciones", []) if c]
    if certs:
        section("Certificaciones")
        for cert in certs:
            story.append(Paragraph(f"• {cert}", s_bullet))

    # Disclaimer
    story.append(Spacer(1, 12))
    story.append(hr(C_DIS, 0.4))
    story.append(Paragraph(DISCLAIMER_TEXT, s_disc))

    doc.build(story)
    buf.seek(0)
    return buf

# ─── Branded PDF builder ──────────────────────────────────────────────────────
LOGO_PATH = "logo.png"
_RL_NAVY  = rl_colors.HexColor("#1B4F8A")
_RL_GOLD  = rl_colors.HexColor("#C8973A")
_RL_LIGHT = rl_colors.HexColor("#8B96A0")
_RL_DARK  = rl_colors.HexColor("#0F1117")

def build_branded_pdf(title: str, content_text: str, person_name: str = "") -> io.BytesIO:
    import os
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=1.5*cm, bottomMargin=2*cm)

    s_hdr  = ParagraphStyle("s_hdr",  fontName="Helvetica-Bold", fontSize=9,
                             textColor=_RL_LIGHT)
    s_title= ParagraphStyle("s_title",fontName="Helvetica-Bold", fontSize=18,
                             textColor=_RL_NAVY, spaceBefore=6, spaceAfter=4)
    s_body = ParagraphStyle("s_body", fontName="Helvetica", fontSize=10.5,
                             textColor=_RL_DARK, leading=16, spaceAfter=6)
    s_bold = ParagraphStyle("s_bold", fontName="Helvetica-Bold", fontSize=10.5,
                             textColor=_RL_DARK, leading=16, spaceAfter=6)
    s_foot = ParagraphStyle("s_foot", fontName="Helvetica", fontSize=8,
                             textColor=_RL_LIGHT, alignment=TA_CENTER)

    story = []

    # Header row: logo | app name | person name
    hdr_left = Paragraph("<b>Analyze-This</b> · CV Optimizer ATS", s_hdr)
    hdr_right= Paragraph(person_name or "", s_hdr)
    if os.path.exists(LOGO_PATH):
        logo = RLImage(LOGO_PATH, width=1.6*cm, height=1.06*cm)
        hdr_data = [[logo, hdr_left, hdr_right]]
        hdr_cols = [2*cm, 9*cm, 6*cm]
    else:
        hdr_data = [[hdr_left, hdr_right]]
        hdr_cols = [11*cm, 6*cm]
    hdr_tbl = Table(hdr_data, colWidths=hdr_cols)
    hdr_tbl.setStyle(TableStyle([
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("ALIGN",  (-1,0),(-1,0), "RIGHT"),
    ]))
    story.append(hdr_tbl)
    story.append(HRFlowable(width="100%", thickness=1.5, color=_RL_NAVY, spaceAfter=10))
    story.append(Paragraph(title, s_title))
    story.append(HRFlowable(width="100%", thickness=0.5, color=_RL_GOLD, spaceAfter=14))

    # Body — parse markdown-lite: **bold**, numbered lists, bullet lines
    import re
    for line in content_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 5))
            continue
        # Escape XML chars
        safe = stripped.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        # Bold lines
        if re.match(r"^\*\*.*\*\*$", stripped):
            story.append(Paragraph(safe.replace("**",""), s_bold))
        # Numbered items
        elif re.match(r"^\d+\.", stripped):
            story.append(Paragraph(f"&nbsp;&nbsp;{safe}", s_body))
        # Bullet items
        elif stripped.startswith("- ") or stripped.startswith("• "):
            story.append(Paragraph(f"• {safe[2:]}", s_body))
        else:
            # Inline **bold**
            safe = re.sub(r"\*\*(.+?)\*\*", r"<b></b>", safe)
            story.append(Paragraph(safe, s_body))

    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="100%", thickness=0.5, color=_RL_LIGHT, spaceAfter=6))
    story.append(Paragraph(
        "Generado por <b>Analyze-This · CV Optimizer ATS</b> · analyze-this-v2.streamlit.app",
        s_foot))
    doc.build(story)
    buf.seek(0)
    return buf

# ─── Next-step tools (carta, entrevista, linkedin) ───────────────────────────
def _run_next_tool(tool: str, nombre: str, titulo: str, resumen: str, skills: str, fn: str, fs):
    api_key = st.session_state.get("user_api_key") or ANTHROPIC_KEY
    job_ctx = st.session_state.get("cv_data", {}).get("titulo_profesional", "este puesto")

    prompts = {
        "carta": f"""Escribe una carta de presentación para el puesto de {titulo}.
Empieza con una idea potente (NO empieces con 'Me postulo para...' ni 'Mi nombre es...').
Conecta la experiencia específica del candidato con las necesidades exactas del puesto.
Termina transmitiendo confianza y con un llamado a la acción natural.
Máximo 200 palabras. Tono profesional pero humano.

Perfil del candidato:
Nombre: {nombre}
Resumen: {resumen}
Habilidades clave: {skills}

Escribe solo la carta, sin títulos ni explicaciones adicionales.""",

        "entrevista": f"""Soy candidato al puesto de {titulo}.
Dame exactamente:
1. Las 8 preguntas más probables en la entrevista para este cargo
2. Para cada pregunta: una estructura de respuesta sólida usando mi experiencia real (método STAR cuando aplique)
3. Al final: 3 preguntas inteligentes que YO le haría al entrevistador para demostrar pensamiento estratégico

Mi perfil:
{resumen}
Habilidades: {skills}

Sé específico y práctico. Evita respuestas genéricas.""",

        "linkedin": f"""Reescribe estas 3 secciones de mi perfil LinkedIn para posicionarme en búsquedas de reclutadores para el puesto de {titulo}:

1. TÍTULO PROFESIONAL (máximo 220 caracteres, incluye keywords del sector)
2. SECCIÓN 'ACERCA DE' (máximo 2.600 caracteres, primera persona, comienza con gancho, termina con CTA)
3. DESCRIPCIÓN DE EXPERIENCIA más reciente (máximo 5 bullets con logros cuantificados)

Mi perfil actual:
{resumen}
Habilidades: {skills}

Haz que cada palabra tenga peso. Optimiza para el algoritmo de LinkedIn y para reclutadores humanos."""
    }

    labels = {
        "carta": "📝 Carta de Presentación",
        "entrevista": "🎤 Preparación de Entrevista",
        "linkedin": "💼 Optimización LinkedIn"
    }

    with st.spinner(f"Generando {labels[tool]}..."):
        try:
            client = anthropic.Anthropic(api_key=api_key)
            msg = client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=2000,
                messages=[{"role": "user", "content": prompts[tool]}]
            )
            result_text = msg.content[0].text.strip()

            st.markdown("---")
            st.subheader(labels[tool])
            st.markdown(result_text)

            # Show inline + PDF download
            st.markdown(result_text)
            title_map = {
                "carta": "Carta de Presentación",
                "entrevista": "Preparación de Entrevista",
                "linkedin": "Optimización LinkedIn"
            }
            try:
                pdf_buf = build_branded_pdf(
                    title_map.get(tool, tool.capitalize()),
                    result_text,
                    nombre
                )
                st.download_button(
                    label="⬇️ Descargar PDF",
                    data=pdf_buf,
                    file_name=f"{tool}_{nombre.replace(' ','_')}.pdf",
                    mime="application/pdf",
                    use_container_width=False
                )
            except Exception as pdf_err:
                st.caption(f"PDF no disponible: {pdf_err}")
                st.download_button(
                    label="⬇️ Descargar .txt",
                    data=result_text.encode("utf-8"),
                    file_name=f"{tool}_{nombre.replace(' ','_')}.txt",
                    mime="text/plain",
                    use_container_width=False
                )
        except Exception as e:
            st.error(f"Error generando {labels[tool]}: {e}")

# ─── Analysis PDF export ──────────────────────────────────────────────────────
def build_analysis_pdf(cv_data: dict) -> io.BytesIO:
    nombre = cv_data.get("nombre", "Candidato")
    titulo = cv_data.get("titulo_profesional", "")
    score  = cv_data.get("score_match", 0)
    ats_ok = cv_data.get("ats_compatible", True)
    ats_det= cv_data.get("ats_detectado", "")
    explain= cv_data.get("score_explicacion", "")
    desglose = cv_data.get("score_desglose", {})
    kw_ok  = cv_data.get("keywords_integradas", [])
    kw_miss= cv_data.get("keywords_faltantes", [])
    coaching = cv_data.get("coaching", [])

    lines = []
    lines.append(f"**Candidato: {nombre}**")
    lines.append(f"Puesto analizado: {titulo}")
    lines.append("")
    lines.append(f"**Score de compatibilidad: {score}%**")
    lines.append(f"ATS compatible: {'Sí' if ats_ok else 'No'}{f'  ·  ATS detectado: {ats_det}' if ats_det else ''}")
    if explain:
        lines.append(explain)
    lines.append("")
    if desglose:
        lines.append("**Desglose del score:**")
        for k, v in desglose.items():
            lines.append(f"- {k.capitalize()}: {v}%")
        lines.append("")
    if kw_ok:
        lines.append(f"**Keywords integradas ({len(kw_ok)}):**")
        lines.append("  " + ", ".join(kw_ok))
        lines.append("")
    if kw_miss:
        lines.append(f"**Keywords ausentes ({len(kw_miss)}):**")
        lines.append("  " + ", ".join(kw_miss))
        lines.append("")
    if coaching:
        lines.append("**Plan de acción:**")
        lines.append("")
        for tip in coaching:
            lines.append(f"**{tip.get('categoria','')}**")
            lines.append(tip.get("tip", ""))
            lines.append("")
    return build_branded_pdf("Análisis de Compatibilidad ATS", "\n".join(lines), nombre)

# ─── Results display ──────────────────────────────────────────────────────────
def show_results(cv_data, fn, fs, max_pages):
    st.markdown("---")
    st.markdown('<div class="section-label">Análisis de Compatibilidad</div>', unsafe_allow_html=True)
    if cv_data.get("_was_truncated"):
        st.markdown('<div class="warn-truncate">ℹ️ CV muy extenso — se analizaron las páginas más recientes. Activa "Cambio de carrera" para procesar el documento completo.</div>', unsafe_allow_html=True)
    ats_detected = cv_data.get("ats_detectado", "")
    ats_ok  = cv_data.get("ats_compatible", True)
    ats_msg = cv_data.get("ats_razon", "")
    score   = cv_data.get("score_match", 0)
    sc_col  = "#22c55e" if score >= 75 else "#f59e0b" if score >= 55 else "#ef4444"
    bc, sc = st.columns([1,2])
    with bc:
        if ats_ok:
            st.markdown('<div class="ats-live"><div class="ats-dot"></div>ATS Compatible</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div style="display:inline-flex;align-items:center;gap:8px;padding:6px 14px;border-radius:99px;background:rgba(239,68,68,0.10);border:1px solid rgba(239,68,68,0.25);font-size:0.82rem;font-weight:600;color:#ef4444">✕ No ATS Compatible</div>', unsafe_allow_html=True)
        if ats_detected:
            st.markdown(f'<div style="margin-top:0.5rem"><span class="at-badge at-badge-blue">🎯 {ats_detected}</span></div>', unsafe_allow_html=True)
        if ats_msg: st.caption(ats_msg)
    with sc:
        st.metric("Match con la oferta (IA)", f"{score}%")
        explain = cv_data.get("score_explicacion","")
        if explain: st.markdown(f'<div class="score-explain">{explain}</div>', unsafe_allow_html=True)
    desglose = cv_data.get("score_desglose",{})
    if desglose:
        with st.expander("📈 Ver desglose"):
            d1,d2,d3,d4=st.columns(4)
            d1.metric("Keywords",    f"{desglose.get('keywords','-')}%")
            d2.metric("Experiencia", f"{desglose.get('experiencia','-')}%")
            d3.metric("Educación",   f"{desglose.get('educacion','-')}%")
            d4.metric("Habilidades", f"{desglose.get('habilidades','-')}%")
    st.markdown("---")
    kw_ok   = cv_data.get("keywords_integradas", [])
    kw_miss = cv_data.get("keywords_faltantes", [])
    if kw_ok or kw_miss:
        chips_ok   = "".join(f'<span class="kw-chip-ok">{k}</span>'   for k in kw_ok)
        chips_miss = "".join(f'<span class="kw-chip-miss">{k}</span>' for k in kw_miss)
        st.markdown(f"""
<div style="display:grid;grid-template-columns:1fr 1fr;gap:1rem;margin-bottom:0.5rem">
  {"" if not kw_ok else f'<div class="kw-section"><span class="kw-label kw-label-ok">✓ Integradas ({len(kw_ok)})</span>{chips_ok}</div>'}
  {"" if not kw_miss else f'<div class="kw-section"><span class="kw-label kw-label-miss">⚠ Ausentes ({len(kw_miss)})</span>{chips_miss}</div>'}
</div>""", unsafe_allow_html=True)
    coaching=cv_data.get("coaching",[])
    if coaching:
        st.markdown("---")
        st.markdown("**Tu plan de acción**")
        st.caption("Recomendaciones para esta postulación — expande cada una:")
        for i, tip in enumerate(coaching):
            cat = tip.get("categoria","")
            tip_txt = tip.get("tip","")
            with st.expander(cat, expanded=(i==0)):
                st.markdown(tip_txt)
    st.markdown("---")
    st.markdown('<div class="section-label">Descarga tu CV optimizado</div>', unsafe_allow_html=True)
    st.markdown('<p style="font-size:0.85rem;color:var(--text2);margin-bottom:1rem">Todos los templates usan el mismo análisis — solo cambia el diseño. Ideal descargar DOCX para editar.</p>', unsafe_allow_html=True)
    nombre = cv_data.get("nombre","cv").replace(" ","_")

    # ── Descarga del análisis completo en PDF ─────────────────────────────
    try:
        analysis_pdf = build_analysis_pdf(cv_data)
        st.download_button(
            label="📄 Descargar análisis completo (PDF)",
            data=analysis_pdf,
            file_name=f"Analisis_ATS_{nombre}.pdf",
            mime="application/pdf",
            use_container_width=False,
            key=f"dl_analysis_{nombre}"
        )
    except Exception:
        pass

    dl1, dl2, dl3, dl4 = st.columns(4)
    for col, (tname, builder) in zip([dl1,dl2,dl3,dl4], BUILDERS.items()):
        info = TEMPLATES.get(tname, {})
        with col:
            st.markdown(f"""<div class="tpl-header">
                <span class="tpl-icon">{info.get('icon','')}</span>
                <div class="tpl-name">{tname}</div>
                <div class="tpl-ideal">{info.get('ideal','')[:22]}</div>
                </div>""", unsafe_allow_html=True)
            try:
                docx_buf = builder(cv_data, fn, float(fs))
                st.download_button(
                    label="⬇️ DOCX",
                    data=docx_buf,
                    file_name=f"CV_ATS_{nombre}_{tname}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key=f"dl_docx_{tname}_{nombre}"
                )
            except Exception as e:
                st.error(f"DOCX: {e}")
            try:
                pdf_buf = build_cv_pdf(cv_data, tname)
                st.download_button(
                    label="⬇️ PDF",
                    data=pdf_buf,
                    file_name=f"CV_ATS_{nombre}_{tname}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    key=f"dl_pdf_{tname}_{nombre}"
                )
            except Exception as e:
                st.error(f"PDF: {e}")

    st.success("✅ ¡Tu CV optimizado está listo!")
    st.caption(f"Tipografía DOCX: {fn} · {fs}pt · {max_pages} página(s)")
    st.markdown(f"""<div class="at-card" style="margin-top:0.6rem;font-size:0.72rem;color:var(--text3);line-height:1.5">
        ⚠️ <strong style="color:var(--text2)">Aviso legal:</strong> {DISCLAIMER_TEXT}
    </div>""", unsafe_allow_html=True)

    # ── Feedback inmediato — calidad del CV generado ──────────────────────
    st.markdown("---")
    with st.expander("⭐ ¿El CV generado se ve bien? (30 segundos)"):
        st.markdown("**¿Qué tan bien quedó el CV?** — Tu feedback mejora la herramienta.")
        st.caption("💡 Para reportar si conseguiste entrevista o trabajo, vuelve al historial arriba cuando tengas novedades.")
        fb_rating = st.select_slider(
            "Calidad del CV generado:",
            options=[1, 2, 3, 4, 5],
            value=5,
            format_func=lambda x: {1:"😞 Muy mal", 2:"😕 Malo", 3:"😐 Aceptable",
                                    4:"😊 Bueno", 5:"🤩 Excelente"}[x]
        )
        fb_comment = st.text_area(
            "Comentario (opcional)",
            placeholder="¿Qué mejorarías del CV generado? ¿Faltó algo? ¿El tono era el correcto?",
            height=80
        )
        if st.button("Enviar opinión", key="btn_feedback"):
            user_data = st.session_state.get("user")
            job = cv_data.get("titulo_profesional", "")
            if save_feedback(user_data.id, user_data.email, fb_rating, fb_comment, job):
                st.success("¡Gracias! 🙏 Recuerda volver al historial para reportar si consigues entrevista.")
            else:
                st.error("Error al enviar. Intenta de nuevo.")

    # ── ¿Qué sigue? — herramientas complementarias ─────────────────────────
    st.markdown("---")
    st.markdown('<div class="section-label">¿Qué sigue?</div>', unsafe_allow_html=True)
    st.markdown('<p style="font-size:0.85rem;color:var(--text2);margin-bottom:1rem">Prepara el resto de tu postulación en segundos, con el mismo CV y oferta:</p>', unsafe_allow_html=True)

    nombre_cv  = cv_data.get("nombre", "")
    titulo_cv  = cv_data.get("titulo_profesional", "")
    resumen_cv = cv_data.get("resumen_profesional", "")
    skills_cv  = ", ".join(cv_data.get("habilidades_tecnicas", [])[:6])

    q1, q2, q3 = st.columns(3)

    with q1:
        st.markdown("""<div class="at-card">
          <div style="font-size:1.1rem;margin-bottom:0.3rem">📝</div>
          <div style="font-size:0.85rem;font-weight:600;color:var(--text);margin-bottom:0.25rem">Carta de presentación</div>
          <div style="font-size:0.75rem;color:var(--text3)">Menos de 200 palabras. Comienza con una idea potente.</div>
        </div>""", unsafe_allow_html=True)
        if st.button("Generar carta", key="btn_carta", use_container_width=True):
            st.session_state["next_tool"] = "carta"
            st.rerun()

    with q2:
        st.markdown("""<div class="at-card">
          <div style="font-size:1.1rem;margin-bottom:0.3rem">🎤</div>
          <div style="font-size:0.85rem;font-weight:600;color:var(--text);margin-bottom:0.25rem">Prep de entrevista</div>
          <div style="font-size:0.75rem;color:var(--text3)">8 preguntas + respuestas en método STAR.</div>
        </div>""", unsafe_allow_html=True)
        if st.button("Preparar entrevista", key="btn_entrevista", use_container_width=True):
            st.session_state["next_tool"] = "entrevista"
            st.rerun()

    with q3:
        st.markdown("""<div class="at-card">
          <div style="font-size:1.1rem;margin-bottom:0.3rem">💼</div>
          <div style="font-size:0.85rem;font-weight:600;color:var(--text);margin-bottom:0.25rem">Optimizar LinkedIn</div>
          <div style="font-size:0.75rem;color:var(--text3)">Título, About y experiencias para reclutadores.</div>
        </div>""", unsafe_allow_html=True)
        if st.button("Optimizar LinkedIn", key="btn_linkedin", use_container_width=True):
            st.session_state["next_tool"] = "linkedin"
            st.rerun()

    # Execute selected tool
    next_tool = st.session_state.get("next_tool")
    if next_tool and cv_data:
        st.session_state.pop("next_tool", None)
        _run_next_tool(next_tool, nombre_cv, titulo_cv, resumen_cv, skills_cv, fn, fs)

# ─── Admin panel ──────────────────────────────────────────────────────────────
def show_admin_panel():
    st.markdown("---")
    with st.expander("🛠️ Panel Admin"):
        stats = get_global_stats()
        a1, a2 = st.columns(2)
        a1.metric("📄 CVs generados (total)", f"{stats['cvs']:,}")
        a2.metric("👥 Usuarios registrados", f"{stats['users']:,}")

        # ── User management ────────────────────────────────────────────────
        # Onboarding stats
        try:
            stages_res = supabase.table("profiles").select("career_stage,how_found").execute()
            if stages_res.data:
                from collections import Counter
                stages = Counter(u.get("career_stage") for u in stages_res.data if u.get("career_stage"))
                found  = Counter(u.get("how_found")    for u in stages_res.data if u.get("how_found"))
                if stages or found:
                    st.markdown("---")
                    st.markdown("**📊 Perfil de usuarios:**")
                    sc1, sc2 = st.columns(2)
                    with sc1:
                        st.caption("Etapa de carrera")
                        for k,v in stages.most_common():
                            st.markdown(f"- {k}: **{v}**")
                    with sc2:
                        st.caption("¿Cómo nos encontraron?")
                        for k,v in found.most_common():
                            st.markdown(f"- {k}: **{v}**")
        except Exception: pass

        st.markdown("---")
        st.markdown("**👤 Gestión de usuarios:**")
        try:
            res = supabase.table("profiles").select("id,email,plan,credits_used_this_month,activation_code").order("created_at", desc=True).limit(20).execute()
            if res.data:
                for u in res.data:
                    uid   = u.get("id")
                    email = u.get("email","")
                    plan  = u.get("plan","free")
                    used  = u.get("credits_used_this_month", 0)
                    limit = PLAN_CREDITS.get(plan, 5)
                    code  = u.get("activation_code","") or ""
                    limit_display = "∞" if plan == "admin" else str(limit)

                    with st.container():
                        uc1, uc2, uc3, uc4 = st.columns([3, 2, 1, 1])
                        with uc1:
                            st.markdown(f"**{email}**")
                            st.caption(f"{used}/{limit_display} análisis · {f'código: {code}' if code else 'sin código'}")
                        with uc2:
                            new_plan = st.selectbox("Plan", ["free","pro_code","pro","admin"],
                                index=["free","pro_code","pro","admin"].index(plan) if plan in ["free","pro_code","pro","admin"] else 0,
                                key=f"plan_{uid}", label_visibility="collapsed")
                        with uc3:
                            if st.button("Cambiar", key=f"change_plan_{uid}"):
                                try:
                                    supabase.table("profiles").update({"plan": new_plan}).eq("id", uid).execute()
                                    st.success("✅"); st.rerun()
                                except: st.error("Error")
                        with uc4:
                            if st.button("↺ Reset", key=f"reset_{uid}", help="Resetear análisis usados este mes"):
                                try:
                                    supabase.table("profiles").update({"credits_used_this_month": 0}).eq("id", uid).execute()
                                    st.success("✅"); st.rerun()
                                except: st.error("Error")
            else:
                st.info("No hay usuarios todavía.")
        except Exception as e:
            st.error(f"Error: {e}")

        # ── Activation codes ───────────────────────────────────────────────
        st.markdown("---")
        st.markdown("**🎟️ Códigos de activación:**")
        codes = get_all_codes()
        if codes:
            for c in codes:
                cid   = c.get("id")
                uses  = c.get("uses_count", 0)
                max_u = c.get("max_uses")
                active = c.get("active", True)
                desc  = c.get("description","")
                grants = c.get("grants_plan","pro_code")
                max_display = str(max_u) if max_u else "∞"

                cc1, cc2, cc3, cc4 = st.columns([2, 2, 1, 1])
                with cc1:
                    st.markdown(f"{'✅' if active else '❌'} **`{c.get('code')}`**")
                    st.caption(f"{desc} · {uses}/{max_display} usos · {grants}")
                with cc2:
                    new_max = st.number_input("Máx usos", min_value=0,
                        value=max_u or 0, key=f"max_{cid}",
                        label_visibility="collapsed",
                        help="0 = ilimitado")
                with cc3:
                    if st.button("Guardar", key=f"save_code_{cid}"):
                        try:
                            supabase.table("activation_codes").update({
                                "max_uses": new_max if new_max > 0 else None
                            }).eq("id", cid).execute()
                            st.success("✅"); st.rerun()
                        except: st.error("Error")
                with cc4:
                    toggle_label = "Desactivar" if active else "Activar"
                    if st.button(toggle_label, key=f"toggle_{cid}"):
                        try:
                            supabase.table("activation_codes").update({"active": not active}).eq("id", cid).execute()
                            st.rerun()
                        except: st.error("Error")
        else:
            st.info("No hay códigos aún.")

        st.markdown("**Crear nuevo código:**")
        nc1, nc2 = st.columns(2)
        with nc1:
            new_code_val  = st.text_input("Código", placeholder="ICI2026", key="new_code")
            new_code_desc = st.text_input("Descripción", placeholder="Ex alumnos ICI UNAB", key="new_desc")
            new_code_plan = st.selectbox("Plan que otorga", ["pro_code","pro"], key="new_plan")
        with nc2:
            new_code_uses = st.number_input("Máx usos (0=ilimitado)", min_value=0, value=0, key="new_uses")
            new_code_exp  = st.text_input("Expira (YYYY-MM-DD, opcional)", key="new_exp")
        if st.button("✅ Crear código", key="btn_create_code"):
            if new_code_val.strip():
                ok = create_code(new_code_val, new_code_desc, int(new_code_uses),
                                 new_code_plan, new_code_exp if new_code_exp else None)
                if ok: st.success(f"Código `{new_code_val.upper()}` creado."); st.rerun()
                else: st.error("Error al crear el código.")
        st.markdown("---")
        st.markdown("**Feedback recibido:**")
        all_fb = get_all_feedback()
        if all_fb:
            for fb in all_fb:
                stars = "⭐" * fb.get("rating", 0)
                approved = fb.get("approved", False)
                fid = fb.get("id")
                comment = fb.get("comment", "—")
                email = fb.get("email", "")
                created = fb.get("created_at", "")[:10]
                job = fb.get("job_title", "")
                status = "✅ Publicado" if approved else "⏳ Pendiente"
                with st.container():
                    st.markdown(f"{stars} **{job}** · {email} · {created} · {status}")
                    st.markdown(f"> {comment}")
                    col_a, col_r, _ = st.columns([1, 1, 4])
                    with col_a:
                        if st.button("✅ Aprobar", key=f"approve_{fid}"):
                            approve_feedback(fid, True); st.rerun()
                    with col_r:
                        if st.button("❌ Rechazar", key=f"reject_{fid}"):
                            approve_feedback(fid, False); st.rerun()
                    st.markdown("---")
        else:
            st.info("No hay feedback todavía.")

# ─── Main app (authenticated) ─────────────────────────────────────────────────
def show_main_app(user, profile):
    plan = profile.get("plan","free")
    credits_left = get_credits_remaining(profile)
    credits_used = profile.get("credits_used_this_month", 0)
    monthly_limit = PLAN_CREDITS.get(plan, 5)

    # ── Sidebar — clean & minimal ──────────────────────────────────────────
    # Smart defaults — what works best for ATS
    max_pages   = 1
    font_family = "Calibri"
    font_size   = 10

    with st.sidebar:
        # ── User identity block ────────────────────────────────────────────
        nombre_sb = profile.get("display_name") or profile.get("email","").split("@")[0]
        initials  = (nombre_sb[0] if nombre_sb else "?").upper()
        plan_pill_class = {"admin": "sb-pill-admin", "pro": "sb-pill-pro"}.get(plan, "sb-pill-free")
        plan_label_map  = {"free": "Free", "pro_code": "Pro", "pro": "Pro", "admin": "Admin"}
        plan_label      = plan_label_map.get(plan, plan.upper())

        st.markdown(f"""
<div class="sb-user">
  <div class="sb-avatar">{initials}</div>
  <div>
    <div class="sb-name">{nombre_sb}</div>
    <div class="sb-plan">
      <span class="sb-pill {plan_pill_class}">{plan_label}</span>
      &nbsp;·&nbsp;{profile.get("email","")[:28]}
    </div>
  </div>
</div>""", unsafe_allow_html=True)

        # ── Credits ───────────────────────────────────────────────────────
        if plan != "admin":
            pct = min(credits_used / monthly_limit, 1.0) if monthly_limit else 0
            bar_class = "zero" if credits_left == 0 else ("low" if credits_left <= 2 else "")
            st.markdown(f"""
<div style="margin-bottom:0.8rem">
  <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:4px">
    <span style="font-size:0.75rem;color:var(--text3);font-weight:500;text-transform:uppercase;letter-spacing:0.05em">Análisis este mes</span>
    <span style="font-size:0.78rem;font-weight:600;color:var(--text2)">{credits_used}<span style="color:var(--text3)">/{monthly_limit}</span></span>
  </div>
  <div class="cred-bar-wrap">
    <div class="cred-bar-fill {bar_class}" style="width:{pct*100:.0f}%"></div>
  </div>
  <div style="font-size:0.75rem;color:var(--text3)">{credits_left} restante{'s' if credits_left != 1 else ''}</div>
</div>""", unsafe_allow_html=True)

            if credits_left == 0:
                st.markdown("""<div style="background:rgba(239,68,68,0.08);border:1px solid rgba(239,68,68,0.2);
                    border-radius:8px;padding:0.6rem 0.8rem;font-size:0.8rem;color:#ef4444;margin-bottom:0.8rem">
                    Sin análisis este mes.<br>
                    <a href="https://ko-fi.com/analyzethis" target="_blank"
                       style="color:#3b82f6;text-decoration:none;font-weight:500">
                    ☕ Apoya en Ko-fi → Pro</a></div>""", unsafe_allow_html=True)
        else:
            st.markdown('<div style="margin-bottom:0.8rem"><span class="at-badge at-badge-blue">∞ Admin · Ilimitado</span></div>', unsafe_allow_html=True)

        if st.session_state.get("api_credits_error"):
            st.warning("Servicio sin saldo.")
            user_key = st.text_input("🔑 Tu API Key", type="password")
            if user_key:
                st.session_state["user_api_key"] = user_key

        st.markdown('<div class="section-label">Mi cuenta</div>', unsafe_allow_html=True)

        # ── Name editor — always visible, not in expander ─────────────────
        current_name = profile.get("display_name") or ""
        new_name = st.text_input("Nombre / apodo", value=current_name,
            placeholder="Ej: Rocío, Juan P.", key="sidebar_display_name",
            help="Como quieres que te llamemos en la app")
        if st.button("Guardar nombre", key="btn_save_name", use_container_width=True):
            if new_name.strip():
                try:
                    supabase.table("profiles").update({
                        "display_name": new_name.strip()
                    }).eq("id", user.id).execute()
                    st.toast("✅ Nombre actualizado")
                    st.rerun()
                except Exception:
                    st.error("Error al guardar")

        # ── Activation code ───────────────────────────────────────────────
        existing_code = profile.get("activation_code", "")
        if existing_code:
            st.markdown(f'<div style="margin-top:0.5rem"><span class="at-badge at-badge-amber">🎟 {existing_code}</span></div>', unsafe_allow_html=True)
        else:
            with st.expander("🎟️ Tengo un código"):
                code_input = st.text_input("Código de activación", key="activate_code_sidebar",
                    placeholder="Ej: ICI2026")
                if st.button("Activar código", key="btn_activate_code", use_container_width=True):
                    if code_input.strip():
                        ok, msg = validate_and_use_code(user.id, code_input)
                        if ok: st.success(msg); st.rerun()
                        else: st.error(msg)

        # ── Regenerate without re-calling Claude ──────────────────────────
        if st.session_state.get("cv_data"):
            st.markdown('<div class="section-label">Formato</div>', unsafe_allow_html=True)
            st.caption("Cambia el diseño sin gastar un análisis.")
            if st.button("🔄 Cambiar template", use_container_width=True):
                st.session_state["regen_docx"] = True

        # ── Advanced format settings ──────────────────────────────────────
        st.markdown('<div class="section-label">Ajustes avanzados</div>', unsafe_allow_html=True)
        with st.expander("⚙️ Formato del CV"):
            st.caption("Por defecto: 1 página, Calibri 10pt — óptimo para ATS.")
            max_pages   = st.slider("Páginas máximas", 1, 5, 1)
            if max_pages > 2:
                st.caption("⚠️ +2 páginas reduce compatibilidad ATS.")
            font_family = st.selectbox("Tipografía",
                ["Calibri","Arial","Georgia","Times New Roman","Trebuchet MS"], index=0)
            font_size   = st.select_slider("Tamaño de letra",
                options=[9, 10, 10.5, 11, 12], value=10)

        st.markdown("---")
        if st.button("🚪 Cerrar sesión", use_container_width=True):
            sign_out()
        st.markdown('<div style="font-size:0.7rem;color:var(--text3);margin-top:0.5rem;text-align:center">Analyze-This · Powered by Claude</div>', unsafe_allow_html=True)

    # ── Header ─────────────────────────────────────────────────────────────
    nombre_usuario = profile.get("display_name") or profile.get("email","").split("@")[0]
    st.markdown(f"""
<div class="app-header">
  <div>
    <div class="app-logo-row">
      <span style="font-size:1.1rem">🎯</span>
      <h1 class="app-title">Analyze-This <span class="app-beta">Beta</span></h1>
    </div>
    <p class="app-subtitle">Hola, <strong style="color:var(--text2)">{nombre_usuario}</strong> — sube tu CV, pega la oferta, descarga listo.</p>
  </div>
</div>""", unsafe_allow_html=True)

    # ── Admin panel ────────────────────────────────────────────────────────
    if plan == "admin":
        show_admin_panel()

    # ── Historial ──────────────────────────────────────────────────────────
    history = get_history(user.id)
    if history:
        # Count pending outcomes
        pending = [h for h in history if not h.get("outcome")]

        # ── Banner de feedback pendiente — visible y directo ───────────────
        if pending and not st.session_state.get("_feedback_banner_dismissed"):
            most_recent = pending[0]
            title_pending = most_recent.get("job_title", "tu última postulación")
            hid_pending   = most_recent.get("id")
            st.markdown(f"""<div class="at-card" style="border-color:rgba(245,158,11,0.25);margin-bottom:1rem">
                <div style="font-size:0.85rem;font-weight:600;color:var(--amber);margin-bottom:0.3rem">
                    🔔 ¿Cómo te fue con <em>{title_pending}</em>?
                </div>
                <div style="font-size:0.78rem;color:var(--text3)">
                    Tu feedback mejora la herramienta. 5 segundos — ¿obtuviste respuesta?
                </div>
            </div>""", unsafe_allow_html=True)

            outcome_labels_q = {
                "got_interview": "🎤 Obtuve entrevista",
                "got_job":       "🎉 Conseguí el trabajo",
                "no_response":   "📭 Sin respuesta",
                "rejected":      "❌ Rechazado",
            }
            fb_cols = st.columns(len(outcome_labels_q) + 1)
            for i, (val, lbl) in enumerate(outcome_labels_q.items()):
                with fb_cols[i]:
                    if st.button(lbl, key=f"quick_outcome_{hid_pending}_{val}",
                                 use_container_width=True):
                        update_outcome(hid_pending, val)
                        st.session_state["_feedback_banner_dismissed"] = True
                        st.toast("¡Gracias por el feedback! 🙏")
                        st.rerun()
            with fb_cols[-1]:
                if st.button("Ahora no", key="dismiss_feedback_banner",
                             use_container_width=True):
                    st.session_state["_feedback_banner_dismissed"] = True
                    st.rerun()

        label = f"📜 Historial ({len(history)})"
        if pending:
            label += f" · 🔔 {len(pending)} pendiente{'s' if len(pending)>1 else ''}"
        with st.expander(label):
            outcome_labels = {
                "got_interview": "🎤 Obtuve entrevista",
                "got_job":       "🎉 Conseguí el trabajo",
                "no_response":   "📭 Sin respuesta",
                "rejected":      "❌ Rechazado",
            }
            for h in history:
                created = h.get("created_at","")[:10]
                score   = h.get("score_match", 0)
                ats     = "✅" if h.get("ats_compatible") else "❌"
                title   = h.get("job_title", "—")
                sc_col  = "🟢" if score >= 75 else "🟡" if score >= 55 else "🔴"
                outcome = h.get("outcome")
                hid     = h.get("id")

                if outcome:
                    outcome_icon = outcome_labels.get(outcome, outcome)
                    st.markdown(f"- {created} · **{title}** · {ats} ATS · {sc_col} {score}% · {outcome_icon}")
                else:
                    col_info, col_select, col_btn = st.columns([3, 2, 1])
                    with col_info:
                        st.markdown(f"**{title}** · {created} · {sc_col} {score}%")
                    with col_select:
                        sel = st.selectbox("¿Resultado?", ["— pendiente —"] + list(outcome_labels.keys()),
                                           key=f"outcome_{hid}",
                                           format_func=lambda x: outcome_labels.get(x, x),
                                           label_visibility="collapsed")
                    with col_btn:
                        if st.button("Guardar", key=f"save_outcome_{hid}"):
                            if sel != "— pendiente —":
                                update_outcome(hid, sel)
                                st.rerun()

    # ── Inputs ─────────────────────────────────────────────────────────────
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("📄 Tu CV")
        career_change = st.toggle(
            "🔄 Cambio de carrera o industria",
            key="career_change_mode",
            help="Activa esto si estás cambiando de área. Se analizará TODA tu experiencia para encontrar habilidades transferibles, sin importar cuán antigua sea."
        )
        if career_change:
            st.info("Modo cambio de carrera activo — se usará toda tu experiencia para encontrar lo transferible al nuevo rol.")
            st.caption("💡 Sube tu historial completo (hasta 30+ páginas). Toda la experiencia cuenta.")
        else:
            st.caption("💡 Sube tu CV completo — se analizarán hasta ~15 páginas priorizando lo más reciente.")
        cv_file = st.file_uploader("Sube tu CV", type=["pdf","docx"], label_visibility="collapsed")
        cv_text_manual = st.text_area("O pega el texto aquí", height=150,
            placeholder="Pega el contenido de tu CV si no tienes archivo...")
        if cv_file and cv_text_manual.strip():
            st.caption("ℹ️ Se usarán tanto el archivo como el texto pegado — el texto se agrega al final del CV.")
    with col2:
        st.subheader("💼 Oferta o Palabras Clave")
        cv_only_mode = st.toggle("🔍 Solo analizar mi CV (sin oferta)",
            key="cv_only_mode",
            help="Analiza tu CV en general: ATS, redacción, modernidad, estructura. Sin comparar con oferta específica.")
        if cv_only_mode:
            st.info("Modo análisis general — recibirás consejos de ATS, redacción y estructura.")
            job_url = ""
            job_description = st.text_area("Área o palabras clave (opcional)", height=100,
                placeholder="Ej: Marketing Digital, Chile, empresas tech. O deja vacío para análisis general.")
        else:
            job_url = st.text_input("🔗 Link de la oferta",
                placeholder="https://www.linkedin.com/jobs/...")
            job_description = st.text_area("O pega el texto aquí", height=175,
                placeholder="Pega el texto de la oferta, o solo palabras clave: 'Gerente proyectos, minería, Excel'")

    st.markdown("---")
    output_lang = st.radio(
        "🌐 Idioma del CV optimizado:",
        options=["es", "en", "pt"],
        format_func=lambda x: {"es": "🇨🇱 Español", "en": "🇺🇸 English", "pt": "🇧🇷 Português"}[x],
        horizontal=True,
        key="output_lang",
        help="El CV original y la oferta pueden estar en idiomas distintos — la IA los compara igual y genera el output en el idioma elegido."
    )
    save_copy = st.toggle(
        "💾 Guardar copia de mi CV para análisis posterior",
        key="save_cv_copy",
        help="Guarda el texto de tu CV original y el CV generado en tu cuenta. Solo tú puedes verlos. Útil para comparar versiones."
    )

    # ── Regenerate only ────────────────────────────────────────────────────
    if st.session_state.get("regen_docx") and st.session_state.get("cv_data"):
        st.session_state["regen_docx"] = False
        show_results(st.session_state["cv_data"], font_family, font_size, max_pages)
        st.stop()

    # ── Main optimize button ───────────────────────────────────────────────
    if credits_left == 0 and plan != "admin" and not st.session_state.get("user_api_key"):
        st.button("🚀 Optimizar mi CV", use_container_width=True, disabled=True)
        st.markdown("""<div class="at-card" style="border-color:rgba(245,158,11,0.2);margin-top:0.5rem">
        <strong style="color:var(--amber)">Sin análisis este mes 🎯</strong><br>
        <span style="font-size:0.83rem;color:var(--text2)">Tu plan Free incluye 10 CVs al mes.<br>
        <a href="mailto:contacto@analyze-this.app" style="color:var(--blue)">Escríbenos para subir a Pro →</a></span>
        </div>""", unsafe_allow_html=True)
        st.stop()

    # Scroll hint if results already exist
    if st.session_state.get("cv_data"):
        st.markdown('<div class="scroll-hint" style="text-align:center;color:var(--text3);font-size:0.78rem;padding:0.3rem">↓ Resultados disponibles abajo</div>', unsafe_allow_html=True)

    if st.button("🚀 Optimizar mi CV", use_container_width=True):
        # Resolve job text — cache scraped result to avoid multiple calls
        final_job = job_description.strip()
        url_key = job_url.strip()
        cached_url = st.session_state.get("_cached_job_url", "")
        cached_text = st.session_state.get("_cached_job_text", "")

        if url_key and is_valid_url(url_key):
            if url_key == cached_url and cached_text:
                final_job = cached_text  # use cache, no re-scrape
            else:
                with st.spinner("🔍 Leyendo la oferta desde el link..."):
                    try:
                        scraped = scrape_job_url(url_key)
                        if scraped:
                            final_job = scraped
                            st.session_state["_cached_job_url"]  = url_key
                            st.session_state["_cached_job_text"] = scraped
                            st.success(f"✅ Oferta leída ({len(scraped):,} caracteres)")
                        else:
                            st.warning("No se pudo extraer texto del link.")
                    except ValueError as e:
                        st.warning(str(e))
        elif url_key:
            st.warning("⚠️ El link no parece válido.")

        if not final_job:
            st.error("⚠️ Pega la oferta o ingresa un link válido.")
            st.stop()

        # Extract CV
        cv_text = ""
        if cv_file:
            with st.spinner("📄 Extrayendo texto del archivo..."):
                try:
                    cv_text = extract_pdf(cv_file) if cv_file.name.lower().endswith(".pdf") \
                              else extract_docx(cv_file)
                except Exception as e:
                    st.error(f"Error al leer el archivo: {e}"); st.stop()

        if cv_text_manual.strip():
            cv_text = cv_text_manual.strip() if not cv_text else cv_text + "\n" + cv_text_manual.strip()

        if not cv_text:
            st.error("⚠️ Sube un CV o pega el texto."); st.stop()

        # Call Claude — show staged progress
        prog = st.progress(0, text="📄 Preparando tu CV para análisis...")
        prog.progress(15, text="🔍 Identificando keywords de la oferta...")
        prog.progress(35, text="🤖 Claude está analizando la compatibilidad...")
        try:
            career_change = st.session_state.get("career_change_mode", False)
            cv_only = st.session_state.get("cv_only_mode", False)
            output_lang = st.session_state.get("output_lang", "es")
            cv_data = optimize_cv(cv_text, final_job, max_pages, font_size, career_change, cv_only, output_lang)
            prog.progress(80, text="✍️ Generando coaching personalizado...")
            st.session_state["cv_data"] = cv_data
            st.session_state["cv_original_text"] = cv_text[:2000]
            st.session_state["api_credits_error"] = False
            prog.progress(100, text="✅ ¡Análisis completado!")
            time.sleep(0.4)
            prog.empty()
            consume_credit(user.id, credits_used)
            history_id = save_history(
                user.id,
                cv_data.get("titulo_profesional", "Rol desconocido"),
                cv_data.get("score_match", 0),
                cv_data.get("ats_compatible", True)
            )
            if history_id:
                st.session_state["last_history_id"] = history_id
                if st.session_state.get("save_cv_copy"):
                    save_cv_copy(user.id, history_id,
                                 cv_text[:5000], cv_data)
        except json.JSONDecodeError:
            prog.empty()
            st.error("Error procesando respuesta. Intenta nuevamente."); st.stop()
        except anthropic.AuthenticationError:
            prog.empty()
            st.error("API Key inválida."); st.stop()
        except anthropic.RateLimitError:
            prog.empty()
            st.session_state["api_credits_error"] = True
            st.error("⚠️ Servicio sin saldo. Ingresa tu API Key en el panel lateral.")
            st.rerun()
        except Exception as e:
            prog.empty()
            st.error(f"Error inesperado: {e}"); st.stop()

        show_results(cv_data, font_family, font_size, max_pages)

    # Show cached results if available (e.g. after sidebar setting change)
    elif st.session_state.get("cv_data"):
        st.info("✅ Análisis previo disponible. Descarga en cualquier formato o haz un nuevo análisis.")
        show_results(st.session_state["cv_data"], font_family, font_size, max_pages)

    st.markdown("---")
    st.markdown('<div style="text-align:center;font-size:0.72rem;color:var(--text3);padding:0.5rem 0">Analyze-This · CV Optimizer ATS · Powered by Claude · Anthropic</div>', unsafe_allow_html=True)

# ─── Guest mode (no login) ────────────────────────────────────────────────────
def show_guest_mode():
    """Let visitor run one full analysis before asking to register."""

    st.markdown(f"""
<div class="app-header">
  <div>
    <div class="app-logo-row">
      <span style="font-size:1.1rem">🎯</span>
      <h1 class="app-title">Analyze-This <span class="app-beta">Beta</span></h1>
    </div>
    <p class="app-subtitle">Sube tu CV y la oferta — ve al instante si eres un buen candidato.</p>
  </div>
</div>""", unsafe_allow_html=True)

    # Show value props
    st.markdown("""
<div class="feat-grid">
  <div class="feat-pill">
    <div class="feat-pill-icon">🎯</div>
    <div class="feat-pill-title">Score ATS real</div>
    <div class="feat-pill-sub">% match con la oferta</div>
  </div>
  <div class="feat-pill">
    <div class="feat-pill-icon">🔑</div>
    <div class="feat-pill-title">Keywords exactas</div>
    <div class="feat-pill-sub">las que el ATS busca</div>
  </div>
  <div class="feat-pill">
    <div class="feat-pill-icon">📄</div>
    <div class="feat-pill-title">CV listo</div>
    <div class="feat-pill-sub">descarga con cuenta gratis</div>
  </div>
</div>
""", unsafe_allow_html=True)

    # ── Inputs ─────────────────────────────────────────────────────────────
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("📄 Tu CV")
        st.caption("💡 Sube tu CV completo — se analizarán hasta ~15 páginas.")
        cv_file = st.file_uploader("Sube tu CV", type=["pdf","docx"], label_visibility="collapsed")
        cv_text_manual = st.text_area("O pega el texto aquí", height=150,
            placeholder="Pega el contenido de tu CV si no tienes archivo...")
    with col2:
        st.subheader("💼 Oferta Laboral")
        guest_cv_only = st.toggle("🔍 Solo analizar mi CV (sin oferta)",
            key="guest_cv_only_mode",
            help="Analiza tu CV en general: ATS, redacción, modernidad. Sin comparar con una oferta específica.")
        if guest_cv_only:
            st.info("Recibirás análisis general de ATS, redacción y estructura.")
            job_url = ""
            job_description = st.text_area("Área o palabras clave (opcional)", height=120,
                placeholder="Ej: Marketing Digital, Chile, fintech. O deja vacío para análisis general.")
        else:
            job_url = st.text_input("🔗 Link de la oferta",
                placeholder="https://www.linkedin.com/jobs/...")
            job_description = st.text_area("O pega el texto aquí", height=160,
                placeholder="Pega el texto completo de la oferta, o solo palabras clave: 'Analista financiero, SAP, Excel, Santiago'")

    # Login button
    _, btn_col = st.columns([4, 1])
    with btn_col:
        if st.button("🔑 Iniciar sesión", use_container_width=True):
            st.session_state["show_auth"] = True
            st.rerun()

    st.markdown("""<div class="at-card" style="font-size:0.83rem;color:var(--text2);margin:0.6rem 0;border-color:var(--blue-bdr)">
        🔓 <strong style="color:var(--text)">Con cuenta gratis obtienes:</strong>
        descarga en 4 templates DOCX · carta de presentación · prep de entrevista ·
        optimización LinkedIn · análisis en PDF · hasta 5 usos al mes.
    </div>""", unsafe_allow_html=True)

    st.markdown("---")

    # If guest already ran analysis, show results with gate — hide inputs
    if st.session_state.get("guest_cv_data"):
        st.info("✅ Análisis completado — regístrate para descargar tu CV optimizado.")
        _show_guest_results(st.session_state["guest_cv_data"])
        return

    if st.button("🔍 Analizar compatibilidad", use_container_width=True, type="primary"):
        # Resolve job
        cv_only_guest = st.session_state.get("guest_cv_only_mode", False)
        final_job = job_description.strip()
        url_key = job_url.strip() if not cv_only_guest else ""
        if url_key and is_valid_url(url_key):
            cached = st.session_state.get("_cached_job_url","")
            if url_key == cached:
                final_job = st.session_state.get("_cached_job_text","") or final_job
            else:
                with st.spinner("🔍 Leyendo oferta..."):
                    try:
                        scraped = scrape_job_url(url_key)
                        if scraped:
                            final_job = scraped
                            st.session_state["_cached_job_url"] = url_key
                            st.session_state["_cached_job_text"] = scraped
                    except Exception: pass

        if not final_job and not cv_only_guest:
            st.error("⚠️ Pega la oferta, un link, o activa 'Solo analizar mi CV'."); st.stop()
        if not final_job:
            final_job = "Análisis general del CV — evaluar ATS, redacción, modernidad y estructura."

        cv_text = ""
        if cv_file:
            with st.spinner("📄 Extrayendo CV..."):
                try:
                    cv_text = extract_pdf(cv_file) if cv_file.name.lower().endswith(".pdf") \
                              else extract_docx(cv_file)
                except Exception as e:
                    st.error(f"Error: {e}"); st.stop()
        if cv_text_manual.strip():
            cv_text = cv_text_manual.strip() if not cv_text else cv_text + "\n" + cv_text_manual.strip()
        if not cv_text:
            st.error("⚠️ Sube tu CV o pega el texto."); st.stop()

        prog = st.progress(0, text="📄 Preparando análisis...")
        prog.progress(20, text="🔍 Leyendo tu CV...")
        prog.progress(40, text="🤖 Analizando compatibilidad...")
        try:
            cv_data = optimize_cv(cv_text, final_job, 1, 10, False, cv_only_guest)
            prog.progress(100, text="✅ ¡Análisis listo!")
            time.sleep(0.4); prog.empty()
            st.session_state["guest_cv_data"] = cv_data
            # Count guest analysis in a lightweight counter table
            try:
                supabase.table("guest_analyses").insert({"created_at": datetime.now(timezone.utc).isoformat()}).execute()
            except Exception:
                pass
            st.rerun()
        except Exception as e:
            prog.empty()
            st.error(f"Error: {e}"); st.stop()

    st.markdown("---")
    st.markdown('<div style="text-align:center;font-size:0.72rem;color:var(--text3);padding:0.5rem 0">Analyze-This · analyze-this-v2.streamlit.app</div>', unsafe_allow_html=True)


def _show_guest_results(cv_data):
    """Show full analysis but gate the download."""
    st.markdown("---")
    st.subheader("📊 Análisis de Compatibilidad")

    ats_detected = cv_data.get("ats_detectado", "")
    ats_ok  = cv_data.get("ats_compatible", True)
    ats_msg = cv_data.get("ats_razon", "")
    score   = cv_data.get("score_match", 0)
    sc_col  = "🟢" if score >= 75 else "🟡" if score >= 55 else "🔴"

    bc, sc = st.columns([1,2])
    with bc:
        if ats_ok: st.success("✅ ATS Compatible")
        else: st.error("❌ No ATS Compatible")
        if ats_detected: st.caption(f"🎯 ATS detectado: **{ats_detected}**")
        if ats_msg: st.caption(ats_msg)
    with sc:
        st.metric(f"{sc_col} Match con la oferta", f"{score}%")
        explain = cv_data.get("score_explicacion","")
        if explain: st.markdown(f'<div class="score-explain">{explain}</div>', unsafe_allow_html=True)

    desglose = cv_data.get("score_desglose",{})
    if desglose:
        with st.expander("📈 Ver desglose"):
            d1,d2,d3,d4=st.columns(4)
            d1.metric("Keywords",    f"{desglose.get('keywords','-')}%")
            d2.metric("Experiencia", f"{desglose.get('experiencia','-')}%")
            d3.metric("Educación",   f"{desglose.get('educacion','-')}%")
            d4.metric("Habilidades", f"{desglose.get('habilidades','-')}%")

    st.markdown("---")
    k1,k2 = st.columns(2)
    with k1:
        kw_ok = cv_data.get("keywords_integradas",[])
        if kw_ok: st.success(f"✅ **Keywords integradas ({len(kw_ok)}):**\n"+", ".join(kw_ok))
    with k2:
        kw_miss = cv_data.get("keywords_faltantes",[])
        if kw_miss: st.warning(f"⚠️ **Keywords ausentes ({len(kw_miss)}):**\n"+", ".join(kw_miss))

    coaching = cv_data.get("coaching",[])
    if coaching:
        st.markdown("---")
        st.markdown("**🎯 Tu Plan de Acción**")
        st.caption("Expande cada recomendación:")
        with st.expander(coaching[0].get("categoria",""), expanded=True):
            st.markdown(coaching[0].get("tip",""))
        if len(coaching) > 1:
            st.caption(f"+ {len(coaching)-1} recomendaciones más — disponibles al registrarte")

    # ── Gate: download requires account ───────────────────────────────────
    st.markdown("---")
    st.markdown("""
<div class="at-card" style="text-align:center;padding:1.8rem;border-color:rgba(245,158,11,0.2)">
  <div style="font-size:1.8rem;margin-bottom:0.6rem">📥</div>
  <div style="font-size:1rem;font-weight:700;color:var(--text);margin-bottom:0.35rem">
    Tu CV optimizado está listo para descargar
  </div>
  <div style="font-size:0.83rem;color:var(--text3);margin-bottom:0">
    Crea tu cuenta gratis y descarga en 4 templates profesionales.<br>
    <span style="color:var(--text2)">Sin tarjeta · 5 análisis/mes incluidos</span>
  </div>
</div>
""", unsafe_allow_html=True)

    col_a, col_b = st.columns(2)
    with col_a:
        if st.button("✨ Crear cuenta gratis", use_container_width=True, type="primary"):
            st.session_state["show_auth"] = True
            st.rerun()
    with col_b:
        if st.button("🔑 Ya tengo cuenta", use_container_width=True):
            st.session_state["show_auth"] = True
            st.rerun()

    # No "analyze another" option — one free analysis per session
    # User must register to continue. This is intentional.
    st.caption("¿Quieres analizar otro CV? Regístrate gratis — tienes 5 análisis/mes.")


# ─── Router ───────────────────────────────────────────────────────────────────
if not supabase:
    st.error("⚠️ Supabase no configurado. Agrega SUPABASE_URL y SUPABASE_KEY en Secrets de Streamlit.")
    st.info("Mientras tanto, usa **app.py** (versión sin usuarios).")
    st.stop()

# Restore persisted session and handle OAuth/magic link callbacks on every rerun
handle_magic_callback()
handle_google_callback()
restore_session()

user = st.session_state.get("user")

if not user:
    if st.session_state.get("show_auth"):
        show_auth_page()
    else:
        show_guest_mode()
else:
    profile = get_profile(user.id)
    if not profile:
        st.error("Error cargando perfil. Intenta cerrar sesión y volver a entrar.")
        if st.button("Cerrar sesión"):
            sign_out()
    else:
        show_main_app(user, profile)


# Setup SQL está en setup.sql — ejecutar en Supabase SQL Editor antes de usar la app.
