"""
CV document builders (DOCX + PDF) — extracted from app2.py, Streamlit-free.
All four templates: Clásico, Moderno, Ejecutivo, Minimalista.
"""
import io
import re


def _x(s) -> str:
    """Escape string for use inside ReportLab XML markup."""
    if not s:
        return ""
    return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors as rl_colors
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
    KeepTogether, Table, TableStyle,
)
from reportlab.platypus import Image as RLImage
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

import os

DISCLAIMER_TEXT = (
    "Este documento fue generado automáticamente por Analyze-This · CV Optimizer ATS "
    "usando inteligencia artificial (Claude, Anthropic). El contenido se basa "
    "exclusivamente en la información provista por el usuario — la herramienta reorganiza "
    "y optimiza, pero no verifica ni valida los datos ingresados. "
    "El usuario es el único responsable de la exactitud de su CV."
)

# ─── DOCX helpers ──────────────────────────────────────────────────────────────

def _section_border(p, border_color):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), border_color)
    pBdr.append(bottom)


def _hdr(doc, title, color_rgb, fn, fs, border_color, prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(fs)
    p.paragraph_format.space_after = Pt(fs * 0.3)
    run = p.add_run(f"{prefix}{title}")
    run.bold = True
    run.font.name = fn
    run.font.size = Pt(fs + 1)
    run.font.color.rgb = RGBColor(*color_rgb)
    _section_border(p, border_color)


def _R(run, fn, fs, bold=False, italic=False, color=None):
    run.font.name = fn
    run.font.size = Pt(float(fs))
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def _exp_block(doc, cv, fn, fs, cargo_color, periodo_color, bullet_prefix, indent_cargo, indent_bullet):
    for exp in cv.get("experiencia", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(fs * 0.6)
        p.paragraph_format.left_indent = Inches(indent_cargo)
        _R(p.add_run(exp.get("cargo", "")), fn, fs, bold=True, color=cargo_color)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(indent_cargo)
        _R(p2.add_run(f"{exp.get('empresa','')}   |   {exp.get('periodo','')}"), fn, fs - 1, italic=True, color=periodo_color)
        for logro in exp.get("logros", []):
            pb = doc.add_paragraph()
            pb.paragraph_format.left_indent = Inches(indent_bullet)
            pb.paragraph_format.space_after = Pt(2)
            _R(pb.add_run(f"{bullet_prefix}{logro}"), fn, fs)


# ─── 4 DOCX templates ──────────────────────────────────────────────────────────

def build_clasico(cv, fn="Calibri", fs=11):
    fs = float(fs)
    doc = DocxDocument()
    s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(1.0)
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)
    DARK = (0x1A, 0x1A, 0x2E); BLUE = (0x2E, 0x75, 0xB6); GRAY = (0x66, 0x66, 0x66)
    def hdr(t): _hdr(doc, t, BLUE, fn, fs, "2E75B6")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _R(p.add_run(cv.get("nombre", "")), fn, fs + 9, bold=True, color=DARK)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _R(p.add_run(cv.get("titulo_profesional", "")), fn, fs + 2, bold=True, color=BLUE)
    parts = [x for x in [cv.get("email"), cv.get("telefono"), cv.get("ubicacion"), cv.get("linkedin")] if x]
    if parts:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _R(p.add_run("  |  ".join(parts)), fn, fs - 1, color=GRAY)
    if cv.get("resumen_profesional"):
        hdr("RESUMEN PROFESIONAL")
        _R(doc.add_paragraph().add_run(cv["resumen_profesional"]), fn, fs)
    if cv.get("experiencia"):
        hdr("EXPERIENCIA PROFESIONAL")
        _exp_block(doc, cv, fn, fs, (0, 0, 0), GRAY, "• ", 0, 0.2)
    if cv.get("educacion"):
        hdr("EDUCACIÓN")
        for edu in cv["educacion"]:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(fs * 0.5)
            _R(p.add_run(edu.get("titulo", "")), fn, fs, bold=True)
            _R(doc.add_paragraph().add_run(f"{edu.get('institucion','')}   |   {edu.get('periodo','')}"), fn, fs - 1, italic=True, color=GRAY)
            if edu.get("detalle"): _R(doc.add_paragraph().add_run(edu["detalle"]), fn, fs - 1)
    if cv.get("habilidades_tecnicas"):
        hdr("HABILIDADES TÉCNICAS"); _R(doc.add_paragraph().add_run("  •  ".join(cv["habilidades_tecnicas"])), fn, fs)
    if cv.get("habilidades_blandas"):
        hdr("COMPETENCIAS"); _R(doc.add_paragraph().add_run("  •  ".join(cv["habilidades_blandas"])), fn, fs)
    if cv.get("idiomas"):
        hdr("IDIOMAS"); _R(doc.add_paragraph().add_run("  |  ".join(cv["idiomas"])), fn, fs)
    certs = [c for c in cv.get("certificaciones", []) if c]
    if certs:
        hdr("CERTIFICACIONES")
        for cert in certs: _R(doc.add_paragraph().add_run(f"• {cert}"), fn, fs)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf


def build_moderno(cv, fn="Calibri", fs=11):
    fs = float(fs)
    doc = DocxDocument()
    s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(0.75)
    s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.8)
    NAVY = (0x1B, 0x4F, 0x72); TEAL = (0x17, 0x8A, 0xCA); GRAY = (0x77, 0x77, 0x77)
    def hdr(t): _hdr(doc, t, NAVY, fn, fs, "17A8CA", prefix="◆  ")
    _R(doc.add_paragraph().add_run(cv.get("nombre", "").upper()), fn, fs + 11, bold=True, color=NAVY)
    _R(doc.add_paragraph().add_run(cv.get("titulo_profesional", "")), fn, fs + 2, bold=True, color=TEAL)
    parts = []
    for icon, key in [("✉", "email"), ("✆", "telefono"), ("⌖", "ubicacion"), ("in", "linkedin")]:
        if cv.get(key): parts.append(f"{icon} {cv[key]}")
    if parts: _R(doc.add_paragraph().add_run("   |   ".join(parts)), fn, fs - 1, color=GRAY)
    p_div = doc.add_paragraph(); p_div.paragraph_format.space_after = Pt(8)
    pPr = p_div._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr"); pPr.append(pBdr)
    btm = OxmlElement("w:bottom")
    for a, v in [("w:val", "single"), ("w:sz", "16"), ("w:space", "1"), ("w:color", "1B4F72")]: btm.set(qn(a), v)
    pBdr.append(btm)
    if cv.get("resumen_profesional"):
        hdr("PERFIL PROFESIONAL")
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.15)
        _R(p.add_run(cv["resumen_profesional"]), fn, fs)
    if cv.get("experiencia"):
        hdr("EXPERIENCIA")
        for exp in cv["experiencia"]:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(fs * 0.7); p.paragraph_format.left_indent = Inches(0.15)
            _R(p.add_run(exp.get("cargo", "")), fn, fs, bold=True, color=NAVY)
            _R(p.add_run("  —  "), fn, fs); _R(p.add_run(exp.get("empresa", "")), fn, fs)
            p2 = doc.add_paragraph(); p2.paragraph_format.left_indent = Inches(0.15)
            _R(p2.add_run(exp.get("periodo", "")), fn, fs - 1, italic=True, color=TEAL)
            for logro in exp.get("logros", []):
                pb = doc.add_paragraph(); pb.paragraph_format.left_indent = Inches(0.35); pb.paragraph_format.space_after = Pt(2)
                _R(pb.add_run(f"▸  {logro}"), fn, fs)
    if cv.get("educacion"):
        hdr("EDUCACIÓN")
        for edu in cv["educacion"]:
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.15); p.paragraph_format.space_before = Pt(fs * 0.5)
            _R(p.add_run(edu.get("titulo", "")), fn, fs, bold=True, color=NAVY)
            _R(p.add_run(f"   |   {edu.get('institucion','')}   |   {edu.get('periodo','')}"), fn, fs - 1)
            if edu.get("detalle"):
                p2 = doc.add_paragraph(); p2.paragraph_format.left_indent = Inches(0.15); _R(p2.add_run(edu["detalle"]), fn, fs - 1)
    if cv.get("habilidades_tecnicas") or cv.get("habilidades_blandas"):
        hdr("HABILIDADES")
        for label, key in [("Técnicas: ", "habilidades_tecnicas"), ("Competencias: ", "habilidades_blandas")]:
            if cv.get(key):
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.15)
                _R(p.add_run(label), fn, fs, bold=True); _R(p.add_run("  •  ".join(cv[key])), fn, fs)
    if cv.get("idiomas"):
        hdr("IDIOMAS"); p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.15)
        _R(p.add_run("  |  ".join(cv["idiomas"])), fn, fs)
    certs = [c for c in cv.get("certificaciones", []) if c]
    if certs:
        hdr("CERTIFICACIONES")
        for cert in certs:
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.15); _R(p.add_run(f"▸  {cert}"), fn, fs)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf


def build_ejecutivo(cv, fn="Calibri", fs=11):
    fs = float(fs)
    sfn = "Georgia" if fn in ["Calibri", "Arial", "Trebuchet MS"] else fn
    doc = DocxDocument()
    s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(1.0)
    s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.8)
    NAVY = (0x1B, 0x2A, 0x4A); GOLD = (0x8B, 0x6C, 0x1E); GRAY = (0x55, 0x55, 0x55)
    def hdr(t): _hdr(doc, t, NAVY, sfn, fs, "1B2A4A")
    p = doc.add_paragraph(); _R(p.add_run(cv.get("nombre", "").upper()), sfn, fs + 10, bold=True, color=NAVY)
    p2 = doc.add_paragraph(); _R(p2.add_run(cv.get("titulo_profesional", "")), sfn, fs + 1, italic=True, color=GOLD)
    parts = [x for x in [cv.get("email"), cv.get("telefono"), cv.get("ubicacion"), cv.get("linkedin")] if x]
    if parts: _R(doc.add_paragraph().add_run("  ·  ".join(parts)), fn, fs - 1, color=GRAY)
    p_div = doc.add_paragraph(); p_div.paragraph_format.space_after = Pt(6)
    pPr = p_div._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr"); pPr.append(pBdr)
    btm = OxmlElement("w:bottom")
    for a, v in [("w:val", "single"), ("w:sz", "24"), ("w:space", "1"), ("w:color", "1B2A4A")]: btm.set(qn(a), v)
    pBdr.append(btm)
    if cv.get("resumen_profesional"):
        hdr("PERFIL EJECUTIVO"); _R(doc.add_paragraph().add_run(cv["resumen_profesional"]), sfn, fs, italic=True)
    if cv.get("experiencia"):
        hdr("TRAYECTORIA PROFESIONAL")
        for exp in cv["experiencia"]:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(fs * 0.7)
            _R(p.add_run(f"■  {exp.get('cargo','')}"), sfn, fs, bold=True, color=NAVY)
            _R(p.add_run(f"  ·  {exp.get('empresa','')}"), sfn, fs)
            p2 = doc.add_paragraph(); p2.paragraph_format.left_indent = Inches(0.25)
            _R(p2.add_run(exp.get("periodo", "")), fn, fs - 1, italic=True, color=GRAY)
            for logro in exp.get("logros", []):
                pb = doc.add_paragraph(); pb.paragraph_format.left_indent = Inches(0.35); pb.paragraph_format.space_after = Pt(2)
                _R(pb.add_run(f"›  {logro}"), sfn, fs)
    if cv.get("educacion"):
        hdr("FORMACIÓN ACADÉMICA")
        for edu in cv["educacion"]:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(fs * 0.5)
            _R(p.add_run(edu.get("titulo", "")), sfn, fs, bold=True, color=NAVY)
            _R(p.add_run(f"  —  {edu.get('institucion','')}  |  {edu.get('periodo','')}"), sfn, fs - 1, color=GRAY)
    all_sk = cv.get("habilidades_tecnicas", []) + cv.get("habilidades_blandas", [])
    if all_sk:
        hdr("COMPETENCIAS"); _R(doc.add_paragraph().add_run("  ■  ".join(all_sk)), sfn, fs)
    if cv.get("idiomas"):
        hdr("IDIOMAS"); _R(doc.add_paragraph().add_run("  |  ".join(cv["idiomas"])), sfn, fs)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf


def build_minimalista(cv, fn="Calibri", fs=11):
    fs = float(fs)
    doc = DocxDocument()
    s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(1.1)
    s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
    DARK = (0x22, 0x22, 0x22); MID = (0x44, 0x44, 0x44); LIGHT = (0x88, 0x88, 0x88)
    def hdr(t):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(fs * 1.2); p.paragraph_format.space_after = Pt(fs * 0.2)
        run = p.add_run(t.upper()); run.bold = True; run.font.name = fn; run.font.size = Pt(fs); run.font.color.rgb = RGBColor(*DARK)
        _section_border(p, "AAAAAA")
    _R(doc.add_paragraph().add_run(cv.get("nombre", "")), fn, fs + 6, bold=True, color=DARK)
    _R(doc.add_paragraph().add_run(cv.get("titulo_profesional", "")), fn, fs + 0.5, color=MID)
    parts = [x for x in [cv.get("email"), cv.get("telefono"), cv.get("ubicacion"), cv.get("linkedin")] if x]
    if parts: _R(doc.add_paragraph().add_run("  |  ".join(parts)), fn, fs - 1, color=LIGHT)
    if cv.get("resumen_profesional"):
        hdr("Resumen"); _R(doc.add_paragraph().add_run(cv["resumen_profesional"]), fn, fs, color=MID)
    if cv.get("experiencia"):
        hdr("Experiencia")
        for exp in cv["experiencia"]:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(fs * 0.5)
            _R(p.add_run(exp.get("cargo", "")), fn, fs, bold=True, color=DARK)
            _R(p.add_run(f"  —  {exp.get('empresa','')}"), fn, fs, color=MID)
            _R(doc.add_paragraph().add_run(exp.get("periodo", "")), fn, fs - 1, italic=True, color=LIGHT)
            for logro in exp.get("logros", []):
                pb = doc.add_paragraph(); pb.paragraph_format.left_indent = Inches(0.2); pb.paragraph_format.space_after = Pt(2)
                _R(pb.add_run(f"- {logro}"), fn, fs, color=MID)
    if cv.get("educacion"):
        hdr("Educación")
        for edu in cv["educacion"]:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(fs * 0.4)
            _R(p.add_run(edu.get("titulo", "")), fn, fs, bold=True, color=DARK)
            _R(p.add_run(f"  —  {edu.get('institucion','')}  ({edu.get('periodo','')})"), fn, fs - 1, color=MID)
    all_sk = cv.get("habilidades_tecnicas", []) + cv.get("habilidades_blandas", [])
    if all_sk:
        hdr("Habilidades"); _R(doc.add_paragraph().add_run("  /  ".join(all_sk)), fn, fs, color=MID)
    if cv.get("idiomas"):
        hdr("Idiomas"); _R(doc.add_paragraph().add_run("  |  ".join(cv["idiomas"])), fn, fs, color=MID)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf


DOCX_BUILDERS = {
    "Clásico": build_clasico,
    "Moderno": build_moderno,
    "Ejecutivo": build_ejecutivo,
    "Minimalista": build_minimalista,
}

TEMPLATES_META = {
    "Clásico":     {"icon": "📋", "color": "#2E75B6", "ideal": "Finanzas · Legal · Gobierno · Roles senior"},
    "Moderno":     {"icon": "✨", "color": "#178ACA", "ideal": "Tech · Startups · Marketing · Diseño"},
    "Ejecutivo":   {"icon": "🏛️", "color": "#1B2A4A", "ideal": "Dirección · C-level · Consultoría · Banca"},
    "Minimalista": {"icon": "⬜", "color": "#444444", "ideal": "Máxima compatibilidad ATS · Cualquier sector"},
}


# ─── PDF builder (reportlab) ───────────────────────────────────────────────────

def build_cv_pdf(cv: dict, template: str = "Clásico") -> io.BytesIO:
    schemes = {
        "Clásico":     {"h": "#2E75B6", "t": "#1A1A2E", "s": "#666666", "acc": "#2E75B6"},
        "Moderno":     {"h": "#1B4F72", "t": "#1B4F72", "s": "#777777", "acc": "#178ACA"},
        "Ejecutivo":   {"h": "#1B2A4A", "t": "#1B2A4A", "s": "#555555", "acc": "#8B6C1E"},
        "Minimalista": {"h": "#222222", "t": "#222222", "s": "#888888", "acc": "#444444"},
    }
    sc = schemes.get(template, schemes["Clásico"])
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm, topMargin=1.8*cm, bottomMargin=2*cm)

    C_H   = rl_colors.HexColor(sc["h"])
    C_T   = rl_colors.HexColor(sc["t"])
    C_S   = rl_colors.HexColor(sc["s"])
    C_ACC = rl_colors.HexColor(sc["acc"])
    C_DIS = rl_colors.HexColor("#AAAAAA")
    fn = "Helvetica"

    def sty(name, **kw):
        base = {"fontName": fn}
        base.update(kw)
        return ParagraphStyle(name, **base)

    s_name    = sty("nm",  fontSize=18, fontName="Helvetica-Bold", textColor=C_T, alignment=TA_CENTER, spaceAfter=3)
    s_title   = sty("tit", fontSize=11, fontName="Helvetica-Bold", textColor=C_H, alignment=TA_CENTER, spaceAfter=3)
    s_contact = sty("con", fontSize=8.5, textColor=C_S, alignment=TA_CENTER, spaceAfter=6)
    s_sec     = sty("sec", fontSize=9.5, fontName="Helvetica-Bold", textColor=C_H, spaceBefore=10, spaceAfter=2)
    s_body    = sty("bod", fontSize=9, textColor=C_T, leading=14, spaceAfter=2)
    s_italic  = sty("ita", fontSize=8.5, textColor=C_S, leading=13, spaceAfter=1)
    s_bullet  = sty("bul", fontSize=9, textColor=C_T, leading=13, leftIndent=12, bulletIndent=0, spaceAfter=1)
    s_disc    = sty("dis", fontSize=7, textColor=C_DIS, leading=10, alignment=TA_CENTER, spaceBefore=8)

    def hr(color=C_H, thickness=0.8):
        return HRFlowable(width="100%", thickness=thickness, color=color, spaceAfter=4)

    story = []
    story.append(Paragraph(_x(cv.get("nombre", "")), s_name))
    if cv.get("titulo_profesional"):
        story.append(Paragraph(_x(cv["titulo_profesional"]), s_title))
    parts = [_x(x) for x in [cv.get("email"), cv.get("telefono"), cv.get("ubicacion"), cv.get("linkedin")] if x]
    if parts:
        story.append(Paragraph("  |  ".join(parts), s_contact))
    story.append(hr())

    def section(title):
        story.append(Paragraph(title.upper(), s_sec))
        story.append(hr(C_ACC, 0.5))

    if cv.get("resumen_profesional"):
        section("Resumen profesional")
        story.append(Paragraph(_x(cv["resumen_profesional"]), s_body))

    if cv.get("experiencia"):
        section("Experiencia profesional")
        for exp in cv["experiencia"]:
            block = []
            block.append(Paragraph(f"<b>{_x(exp.get('cargo',''))}</b>  —  {_x(exp.get('empresa',''))}", s_body))
            block.append(Paragraph(_x(exp.get("periodo", "")), s_italic))
            for logro in exp.get("logros", []):
                block.append(Paragraph(f"• {_x(logro)}", s_bullet))
            story.append(KeepTogether(block))
            story.append(Spacer(1, 4))

    if cv.get("educacion"):
        section("Educación")
        for edu in cv["educacion"]:
            story.append(Paragraph(
                f"<b>{_x(edu.get('titulo',''))}</b>  —  {_x(edu.get('institucion',''))}  ({_x(edu.get('periodo',''))})", s_body))
            if edu.get("detalle"):
                story.append(Paragraph(_x(edu["detalle"]), s_italic))

    sk_tec = cv.get("habilidades_tecnicas", [])
    sk_bla = cv.get("habilidades_blandas", [])
    if sk_tec or sk_bla:
        section("Habilidades")
        if sk_tec:
            story.append(Paragraph("<b>Técnicas:</b>  " + "  ·  ".join(_x(s) for s in sk_tec), s_body))
        if sk_bla:
            story.append(Paragraph("<b>Competencias:</b>  " + "  ·  ".join(_x(s) for s in sk_bla), s_body))

    if cv.get("idiomas"):
        section("Idiomas")
        story.append(Paragraph("  |  ".join(_x(i) for i in cv["idiomas"]), s_body))

    certs = [c for c in cv.get("certificaciones", []) if c]
    if certs:
        section("Certificaciones")
        for cert in certs:
            story.append(Paragraph(f"• {_x(cert)}", s_bullet))

    story.append(Spacer(1, 12))
    story.append(hr(C_DIS, 0.4))
    story.append(Paragraph(DISCLAIMER_TEXT, s_disc))
    doc.build(story)
    buf.seek(0)
    return buf


# ─── Branded PDF (for tools: carta, entrevista, linkedin + analysis) ───────────

_LOGO_PATH = os.path.join(os.path.dirname(__file__), "..", "static", "logo.png")
_RL_NAVY  = rl_colors.HexColor("#1B4F8A")
_RL_GOLD  = rl_colors.HexColor("#C8973A")
_RL_LIGHT = rl_colors.HexColor("#8B96A0")
_RL_DARK  = rl_colors.HexColor("#0F1117")


def build_branded_pdf(title: str, content_text: str, person_name: str = "") -> io.BytesIO:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=1.5*cm, bottomMargin=2*cm)

    s_hdr   = ParagraphStyle("s_hdr",  fontName="Helvetica-Bold", fontSize=9,  textColor=_RL_LIGHT)
    s_title = ParagraphStyle("s_title", fontName="Helvetica-Bold", fontSize=18, textColor=_RL_NAVY, spaceBefore=6, spaceAfter=4)
    s_body  = ParagraphStyle("s_body",  fontName="Helvetica", fontSize=10.5, textColor=_RL_DARK, leading=16, spaceAfter=6)
    s_bold  = ParagraphStyle("s_bold",  fontName="Helvetica-Bold", fontSize=10.5, textColor=_RL_DARK, leading=16, spaceAfter=6)
    s_foot  = ParagraphStyle("s_foot",  fontName="Helvetica", fontSize=8, textColor=_RL_LIGHT, alignment=TA_CENTER)

    story = []

    hdr_left  = Paragraph("<b>Analyze-This</b> · CV Optimizer ATS", s_hdr)
    hdr_right = Paragraph(person_name or "", s_hdr)
    logo_path = os.path.normpath(_LOGO_PATH)
    if os.path.exists(logo_path):
        logo = RLImage(logo_path, width=1.6*cm, height=1.06*cm)
        hdr_data = [[logo, hdr_left, hdr_right]]
        hdr_cols = [2*cm, 9*cm, 6*cm]
    else:
        hdr_data = [[hdr_left, hdr_right]]
        hdr_cols = [11*cm, 6*cm]
    hdr_tbl = Table(hdr_data, colWidths=hdr_cols)
    hdr_tbl.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN",  (-1, 0), (-1, 0), "RIGHT"),
    ]))
    story.append(hdr_tbl)
    story.append(HRFlowable(width="100%", thickness=1.5, color=_RL_NAVY, spaceAfter=10))
    story.append(Paragraph(title, s_title))
    story.append(HRFlowable(width="100%", thickness=0.5, color=_RL_GOLD, spaceAfter=14))

    for line in content_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 5))
            continue
        safe = stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if re.match(r"^\*\*.*\*\*$", stripped):
            story.append(Paragraph(safe.replace("**", ""), s_bold))
        elif re.match(r"^\d+\.", stripped):
            story.append(Paragraph(f"&nbsp;&nbsp;{safe}", s_body))
        elif stripped.startswith("- ") or stripped.startswith("• "):
            story.append(Paragraph(f"• {safe[2:]}", s_body))
        else:
            safe = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", safe)
            story.append(Paragraph(safe, s_body))

    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="100%", thickness=0.5, color=_RL_LIGHT, spaceAfter=6))
    story.append(Paragraph(
        "Generado por <b>Analyze-This · CV Optimizer ATS</b> · analyze-this-production.up.railway.app",
        s_foot))
    doc.build(story)
    buf.seek(0)
    return buf


def build_analysis_pdf(cv_data: dict) -> io.BytesIO:
    nombre  = cv_data.get("nombre", "Candidato")
    titulo  = cv_data.get("titulo_profesional", "")
    score   = cv_data.get("score_match", 0)
    ats_ok  = cv_data.get("ats_compatible", True)
    ats_det = cv_data.get("ats_detectado", "")
    explain = cv_data.get("score_explicacion", "")
    desglose  = cv_data.get("score_desglose", {})
    kw_ok   = cv_data.get("keywords_integradas", [])
    kw_miss = cv_data.get("keywords_faltantes", [])
    coaching = cv_data.get("coaching", [])

    lines = []
    lines.append(f"**Candidato: {nombre}**")
    lines.append(f"Puesto analizado: {titulo}")
    lines.append("")
    lines.append(f"**Score de compatibilidad: {score}%**")
    lines.append(f"ATS compatible: {'Sí' if ats_ok else 'No'}{f'  ·  ATS detectado: {ats_det}' if ats_det else ''}")
    if explain:
        lines.append(explain)
    lines.append("")
    if desglose:
        lines.append("**Desglose del score:**")
        for k, v in desglose.items():
            lines.append(f"- {k.capitalize()}: {v}%")
        lines.append("")
    if kw_ok:
        lines.append(f"**Keywords integradas ({len(kw_ok)}):**")
        lines.append("  " + ", ".join(kw_ok))
        lines.append("")
    if kw_miss:
        lines.append(f"**Keywords ausentes ({len(kw_miss)}):**")
        lines.append("  " + ", ".join(kw_miss))
        lines.append("")
    if coaching:
        lines.append("**Plan de acción:**")
        lines.append("")
        for tip in coaching:
            lines.append(f"**{tip.get('categoria', '')}**")
            lines.append(tip.get("tip", ""))
            lines.append("")
    return build_branded_pdf("Análisis de Compatibilidad ATS", "\n".join(lines), nombre)
